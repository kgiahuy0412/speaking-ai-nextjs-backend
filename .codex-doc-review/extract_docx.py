from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree


def iter_blocks(parent):
    if isinstance(parent, DocumentType):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def extract_xml_text(zf: zipfile.ZipFile, part: str) -> list[str]:
    if part not in zf.namelist():
        return []
    root = etree.fromstring(zf.read(part))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    items = []
    for p in root.xpath(".//w:p", namespaces=ns):
        text = "".join(p.xpath(".//w:t/text()", namespaces=ns)).strip()
        if text:
            items.append(text)
    return items


def extract_last_rendered_pages(zf: zipfile.ZipFile) -> dict:
    document_part = "word/document.xml"
    if document_part not in zf.namelist():
        return {"marker_count": 0, "declared_pages": None, "pages": []}

    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    root = etree.fromstring(zf.read(document_part))
    rendered_tag = f"{{{w_ns}}}lastRenderedPageBreak"
    rendered_count = len(root.findall(f".//{rendered_tag}"))
    marker_tag = rendered_tag if rendered_count else None
    pages: list[list[str]] = [[]]

    def append(value: str) -> None:
        if value:
            pages[-1].append(value)

    def walk(node) -> None:
        if marker_tag and node.tag == marker_tag:
            pages.append([])
            return
        local = etree.QName(node).localname
        if local == "t":
            append(node.text or "")
            return
        if local == "tab":
            append("\t")
            return
        if local in {"cr", "br"}:
            append("\n")
            return
        for child in node:
            walk(child)
        if local in {"p", "tr"}:
            append("\n")
        elif local == "tc":
            append("\t")

    walk(root)
    page_text = ["\n".join(line.strip() for line in "".join(parts).splitlines() if line.strip()) for parts in pages]

    declared_pages = None
    app_part = "docProps/app.xml"
    if app_part in zf.namelist():
        app_root = etree.fromstring(zf.read(app_part))
        values = app_root.xpath("//*[local-name()='Pages']/text()")
        if values:
            try:
                declared_pages = int(values[0])
            except ValueError:
                declared_pages = None

    return {
        "marker_count": rendered_count,
        "declared_pages": declared_pages,
        "pages": page_text,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input")
    parser.add_argument("output")
    args = parser.parse_args()

    src = Path(args.input)
    doc = Document(src)
    blocks = []
    p_index = 0
    t_index = 0
    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            p_index += 1
            text = block.text.strip()
            if text:
                blocks.append(
                    {
                        "type": "paragraph",
                        "index": p_index,
                        "style": block.style.name if block.style else "",
                        "text": text,
                    }
                )
        else:
            t_index += 1
            rows = []
            for row in block.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            blocks.append({"type": "table", "index": t_index, "rows": rows})

    headers = []
    footers = []
    for section_no, section in enumerate(doc.sections, start=1):
        header_text = [p.text.strip() for p in section.header.paragraphs if p.text.strip()]
        footer_text = [p.text.strip() for p in section.footer.paragraphs if p.text.strip()]
        if header_text:
            headers.append({"section": section_no, "paragraphs": header_text})
        if footer_text:
            footers.append({"section": section_no, "paragraphs": footer_text})

    with zipfile.ZipFile(src) as zf:
        extras = {
            "footnotes": extract_xml_text(zf, "word/footnotes.xml"),
            "endnotes": extract_xml_text(zf, "word/endnotes.xml"),
            "comments": extract_xml_text(zf, "word/comments.xml"),
            "last_rendered_pagination": extract_last_rendered_pages(zf),
        }

    result = {
        "source": str(src),
        "paragraph_count": p_index,
        "table_count": t_index,
        "blocks": blocks,
        "headers": headers,
        "footers": footers,
        **extras,
    }
    Path(args.output).write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
