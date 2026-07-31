from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


path = Path("reports/bao-cao-ai-giong-noi-v2-chi-tiet-model-va-chi-phi.docx")
assert path.exists() and path.stat().st_size > 100_000

with ZipFile(path) as archive:
    assert archive.testzip() is None
    xml = archive.read("word/document.xml")

root = etree.fromstring(xml)
ns = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}

doc = Document(path)
full_text = "\n".join(p.text for p in doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            full_text += "\n" + cell.text

for required in [
    "8. Cách làm, triển khai và chi phí theo từng model",
    "8.1 openai/whisper-large-v3-turbo",
    "8.2 vinai/PhoWhisper-small",
    "8.3 vinai/PhoWhisper-medium",
    "8.4 Cloudflare Whisper Turbo",
    "14. Nguồn tham khảo và giả định",
    "~75-140 triệu VND",
    "~90-180 triệu VND",
]:
    assert required in full_text, required

for forbidden in ["TODO", "PLACEHOLDER", "turn11search", "cite"]:
    assert forbidden not in full_text

for table in doc.tables:
    tbl_pr = table._tbl.tblPr
    assert tbl_pr.find(qn("w:tblW")) is not None
    assert tbl_pr.find(qn("w:tblInd")) is not None
    assert table._tbl.tblGrid is not None
    grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid.gridCol_lst]
    assert sum(grid_widths) == 9360, sum(grid_widths)
    if len(table.columns) > 1:
        header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
        assert header_tr_pr.find(qn("w:tblHeader")) is not None

descriptions = root.xpath("//wp:docPr/@descr", namespaces=ns)
assert len([value for value in descriptions if value]) == len(doc.inline_shapes) == 2

heading_texts = [
    p.text for p in doc.paragraphs
    if p.style and p.style.name in {"Heading 1", "Heading 2", "Heading 3"}
]
top_headings = [text for text in heading_texts if text[:1].isdigit() and ". " in text]
assert len(top_headings) >= 14

hyperlinks = root.xpath("//w:hyperlink", namespaces=ns)
assert len(hyperlinks) >= 19

print({
    "path": str(path.resolve()),
    "bytes": path.stat().st_size,
    "paragraphs": len(doc.paragraphs),
    "tables": len(doc.tables),
    "headings": len(heading_texts),
    "top_level_headings": len(top_headings),
    "inline_images_with_alt": len(descriptions),
    "hyperlinks": len(hyperlinks),
    "sections": len(doc.sections),
})
