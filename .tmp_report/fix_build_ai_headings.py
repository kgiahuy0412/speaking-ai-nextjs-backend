from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path(
    r"D:\Code\HuaMei\App_noi\be\3_23th7\speaking-ai-nextjs-backend\reports\Build AI - ban hoan thien.docx"
)
OUTPUT = Path(
    r"D:\Code\HuaMei\App_noi\be\3_23th7\speaking-ai-nextjs-backend\reports\Build AI - sua tieu de muc.docx"
)


HEADING_REPLACEMENTS = {
    "Phương án 1 - Cloudflare Workers Paid + Whisper":
        "3. Phương án 1 - Cloudflare Workers Paid + Whisper",
}

EXPECTED_HEADING_TEXTS = {
    "3. Phương án 1 - Cloudflare Workers Paid + Whisper": "Heading 1",
    "3.1. Kịch bản 10 phút/người/tháng": "Heading 2",
    "3.2. Kịch bản 10 phút/người/ngày": "Heading 2",
    "3.3. Ưu điểm và giới hạn": "Heading 2",
    "4. Phương án 2 - Fine-tune PhoWhisper-small": "Heading 1",
    "4.1. Cấu hình GPU và chi phí huấn luyện": "Heading 2",
    "4.2. Chi phí triển khai model sau huấn luyện": "Heading 2",
    "5. So sánh": "Heading 1",
}


def set_style_font(style, name: str, size: float, color: str, bold: bool):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)

    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def set_keep_with_next(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    keep_next = ppr.find(qn("w:keepNext"))
    if keep_next is None:
        keep_next = OxmlElement("w:keepNext")
        ppr.append(keep_next)

    keep_lines = ppr.find(qn("w:keepLines"))
    if keep_lines is None:
        keep_lines = OxmlElement("w:keepLines")
        ppr.append(keep_lines)


def replace_paragraph_text_preserving_properties(paragraph, new_text: str):
    paragraph_properties = deepcopy(paragraph._p.pPr)
    for child in list(paragraph._p):
        paragraph._p.remove(child)
    if paragraph_properties is not None:
        paragraph._p.append(paragraph_properties)
    paragraph.add_run(new_text)


def all_text_items(document):
    items = []
    for paragraph in document.paragraphs:
        items.append(("paragraph", paragraph.text))
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                items.append(
                    (
                        f"table:{table_index}:{row_index}:{cell_index}",
                        "\n".join(paragraph.text for paragraph in cell.paragraphs),
                    )
                )
    return items


def main():
    before = Document(SOURCE)
    before_items = all_text_items(before)
    before_counts = (
        len(before.paragraphs),
        len(before.tables),
        tuple((len(table.rows), len(table.columns)) for table in before.tables),
    )

    document = Document(SOURCE)

    heading_1 = document.styles["Heading 1"]
    set_style_font(heading_1, "Calibri", 16, "1F4E79", True)
    heading_1.paragraph_format.space_before = Pt(16)
    heading_1.paragraph_format.space_after = Pt(7)
    heading_1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    heading_1.paragraph_format.keep_with_next = True
    heading_1.paragraph_format.keep_together = True

    heading_2 = document.styles["Heading 2"]
    set_style_font(heading_2, "Calibri", 13, "2E74B5", True)
    heading_2.paragraph_format.space_before = Pt(12)
    heading_2.paragraph_format.space_after = Pt(5)
    heading_2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    heading_2.paragraph_format.keep_with_next = True
    heading_2.paragraph_format.keep_together = True

    changed_titles = 0
    styled_headings = 0
    for paragraph in document.paragraphs:
        original_text = paragraph.text.strip()
        if original_text in HEADING_REPLACEMENTS:
            replace_paragraph_text_preserving_properties(
                paragraph, HEADING_REPLACEMENTS[original_text]
            )
            changed_titles += 1

        target_style = EXPECTED_HEADING_TEXTS.get(paragraph.text.strip())
        if target_style:
            paragraph.style = document.styles[target_style]
            set_keep_with_next(paragraph)
            styled_headings += 1

    if changed_titles != len(HEADING_REPLACEMENTS):
        raise RuntimeError(
            f"Expected {len(HEADING_REPLACEMENTS)} heading text correction, got {changed_titles}"
        )
    if styled_headings != len(EXPECTED_HEADING_TEXTS):
        raise RuntimeError(
            f"Expected {len(EXPECTED_HEADING_TEXTS)} styled headings, got {styled_headings}"
        )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)

    after = Document(OUTPUT)
    after_items = all_text_items(after)
    after_counts = (
        len(after.paragraphs),
        len(after.tables),
        tuple((len(table.rows), len(table.columns)) for table in after.tables),
    )

    expected_after_items = []
    for kind, text in before_items:
        replacement = HEADING_REPLACEMENTS.get(text.strip()) if kind == "paragraph" else None
        expected_after_items.append((kind, replacement if replacement is not None else text))

    if before_counts != after_counts:
        raise RuntimeError(f"Structure changed: before={before_counts}, after={after_counts}")
    if expected_after_items != after_items:
        differences = [
            (index, expected, actual)
            for index, (expected, actual) in enumerate(
                zip(expected_after_items, after_items)
            )
            if expected != actual
        ]
        raise RuntimeError(
            "Unexpected text change outside the approved heading correction: "
            + repr(differences[:5])
        )

    actual_heading_map = {
        paragraph.text.strip(): paragraph.style.name
        for paragraph in after.paragraphs
        if paragraph.text.strip() in EXPECTED_HEADING_TEXTS
    }
    if actual_heading_map != EXPECTED_HEADING_TEXTS:
        raise RuntimeError(
            f"Heading hierarchy mismatch: expected={EXPECTED_HEADING_TEXTS}, actual={actual_heading_map}"
        )

    print(f"OUTPUT={OUTPUT}")
    print(f"STRUCTURE_UNCHANGED=YES {after_counts}")
    print("BODY_AND_TABLE_TEXT_UNCHANGED=YES")
    print(f"HEADING_TEXT_CORRECTIONS={changed_titles}")
    print(f"HEADINGS_STANDARDIZED={styled_headings}")


if __name__ == "__main__":
    main()
