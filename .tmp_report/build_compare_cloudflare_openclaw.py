from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import build_asr_report as base


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
WORK_DIR = ROOT / ".tmp_report" / "compare_cloudflare_openclaw"
OUT_PATH = OUT_DIR / "so-sanh-cloudflare-va-openclaw-openrouter-10000-nguoi-dung.docx"
CHART_PATH = WORK_DIR / "so-sanh-chi-phi-100000-phut.png"

USD_TO_VND = 26_000
USERS = 10_000
MINUTES_PER_USER = 10
TOTAL_MINUTES = USERS * MINUTES_PER_USER
DAYS = 30

CF_RATE_PER_MIN = 0.00051
CF_WORKERS_PAID = 5.0
CF_NEURON_RATE = 0.011 / 1_000
CF_FREE_NEURONS_PER_DAY = 10_000
CF_FREE_VALUE_PER_DAY = CF_FREE_NEURONS_PER_DAY * CF_NEURON_RATE
CF_FREE_MINUTES_PER_DAY = CF_FREE_VALUE_PER_DAY / CF_RATE_PER_MIN
CF_MONTHLY_FREE_MINUTES = CF_FREE_MINUTES_PER_DAY * DAYS
CF_GROSS_AI = TOTAL_MINUTES * CF_RATE_PER_MIN
CF_AI_AFTER_FREE = max(0.0, CF_GROSS_AI - CF_FREE_VALUE_PER_DAY * DAYS)
CF_TOTAL = CF_WORKERS_PAID + CF_AI_AFTER_FREE
CF_ONE_DAY_TOTAL = CF_WORKERS_PAID + max(0.0, CF_GROSS_AI - CF_FREE_VALUE_PER_DAY)

OR_RATE_PER_HOUR = 0.04
OR_RATE_PER_MIN = OR_RATE_PER_HOUR / 60
OR_MODEL_COST = TOTAL_MINUTES * OR_RATE_PER_MIN
OPENCLAW_VPS_LOW = 5.0
OPENCLAW_VPS_HIGH = 20.0
OR_TOTAL_LOW = OR_MODEL_COST + OPENCLAW_VPS_LOW
OR_TOTAL_HIGH = OR_MODEL_COST + OPENCLAW_VPS_HIGH


def money(value: float) -> str:
    return f"{value:,.2f} USD".replace(",", " ")


def vnd(value: float) -> str:
    return f"{value * USD_TO_VND:,.0f} VNĐ".replace(",", ".")


def set_section_chrome(doc: Document):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        header = section.header
        hp = header.paragraphs[0]
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        base.set_run_font(
            hp.add_run("SO SÁNH PHƯƠNG ÁN ASR  |  10.000 NGƯỜI DÙNG"),
            size=8.5,
            color=base.MUTED,
            bold=True,
        )

        first_header = section.first_page_header
        fp = first_header.paragraphs[0]
        fp.clear()
        base.set_run_font(
            fp.add_run("BÁO CÁO RA QUYẾT ĐỊNH  |  CẬP NHẬT 27.07.2026"),
            size=8.5,
            color=base.MUTED,
            bold=True,
        )

        footer = section.footer
        p = footer.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        base.set_run_font(p.add_run("Trang "), size=9, color=base.MUTED)
        base.add_page_field(p, "PAGE")
        base.set_run_font(p.add_run(" / "), size=9, color=base.MUTED)
        base.add_page_field(p, "NUMPAGES")

        first_footer = section.first_page_footer
        ffp = first_footer.paragraphs[0]
        ffp.clear()
        ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        base.set_run_font(
            ffp.add_run("Ước tính phục vụ lập ngân sách; chưa gồm thuế và các dịch vụ ngoài phạm vi nêu trong báo cáo."),
            size=8.5,
            color=base.MUTED,
            italic=True,
        )


def add_text(doc: Document, text: str, *, bold_lead: str | None = None, style: str = "Normal"):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        base.set_run_font(p.add_run(bold_lead), bold=True)
        base.set_run_font(p.add_run(text[len(bold_lead):]))
    else:
        base.set_run_font(p.add_run(text))
    return p


def add_source(doc: Document, title: str, url: str):
    p = doc.add_paragraph(style="Normal")
    base.apply_num(p, doc._decimal_num_id)
    p.paragraph_format.space_after = Pt(6)
    base.set_run_font(p.add_run(f"{title}: "), size=9.5)
    base.add_hyperlink(p, "Mở tài liệu chính thức", url)


def add_callout(doc: Document, label: str, text: str, *, kind: str):
    table = base.add_callout(doc, label, text, kind=kind)
    # Re-apply the preset's exact fixed geometry and 120-DXA cell margins.
    base.set_table_geometry(table, [9360], indent_dxa=120)
    base.set_repeat_table_header(table.rows[0])
    return table


def make_cost_chart(path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1600, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    regular_path = Path("C:/Windows/Fonts/arial.ttf")
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    font = lambda size, bold=False: ImageFont.truetype(str(bold_path if bold else regular_path), size)

    draw.text((90, 60), "Chi phí ước tính cho 100.000 phút audio / tháng", font=font(48, True), fill="#17365D")
    draw.text((90, 125), "Giả định tải phân bổ đều trong 30 ngày; OpenClaw chạy trên VPS riêng", font=font(25), fill="#667085")

    items = [
        ("Cloudflare", CF_TOTAL, "đã gồm Workers Paid"),
        ("OpenClaw + OpenRouter", OR_TOTAL_LOW, "VPS 5 USD/tháng"),
        ("OpenClaw + OpenRouter", OR_TOTAL_HIGH, "VPS 20 USD/tháng"),
    ]
    colors = ["#2E74B5", "#8A6500", "#9B1C1C"]
    x0, y0, chart_w, bar_h, gap = 480, 245, 930, 115, 105
    max_val = 90
    for idx, ((label, value, note), color) in enumerate(zip(items, colors)):
        y = y0 + idx * (bar_h + gap)
        draw.text((90, y + 14), label, font=font(28, True), fill="#0B2545")
        draw.text((90, y + 55), note, font=font(21), fill="#667085")
        draw.rounded_rectangle((x0, y, x0 + chart_w, y + bar_h), radius=18, fill="#F2F4F7")
        width = int(chart_w * value / max_val)
        draw.rounded_rectangle((x0, y, x0 + width, y + bar_h), radius=18, fill=color)
        draw.text((x0 + 20, y + 27), f"{value:.2f} USD", font=font(32, True), fill="#FFFFFF")

    draw.text((90, 805), "Nguồn giá model: Cloudflare và OpenRouter. Chi phí VPS là khoảng lập ngân sách, không phải báo giá nhà cung cấp.", font=font(20), fill="#667085")
    img.save(path, quality=95)


def cover(doc: Document):
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph(style="Kicker")
    base.set_run_font(p.add_run("BÁO CÁO SO SÁNH CHI PHÍ VÀ TRIỂN KHAI"), size=10, color=base.GOLD, bold=True)

    p = doc.add_paragraph(style="Report Title")
    p.paragraph_format.space_after = Pt(8)
    base.set_run_font(p.add_run("Cloudflare hay OpenClaw + OpenRouter?"), size=28, color=base.NAVY, bold=True)

    p = doc.add_paragraph(style="Report Subtitle")
    base.set_run_font(p.add_run("Kịch bản 10.000 người dùng, mỗi người 10 phút audio"), size=15, color=base.DARK_BLUE)

    metadata = [
        ("Khối lượng", "100.000 phút audio"),
        ("Kỳ tính chính", "Một tháng, tải phân bổ đều trong 30 ngày"),
        ("Tỷ giá kế hoạch", "1 USD = 26.000 VNĐ"),
        ("Mục tiêu", "Tối ưu chi phí nhưng giữ chất lượng nhận dạng tiếng Việt"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        base.set_run_font(p.add_run(f"{label}: "), size=10.5, bold=True, color=base.INK)
        base.set_run_font(p.add_run(value), size=10.5, color=base.INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_callout(
        doc,
        "Kết luận",
        f"Giữ Cloudflare là phương án chính. Tổng ước tính {money(CF_TOTAL)}/tháng, thấp hơn OpenClaw + OpenRouter khoảng {money(OR_TOTAL_LOW - CF_TOTAL)} đến {money(OR_TOTAL_HIGH - CF_TOTAL)}/tháng trong cùng kịch bản.",
        kind="positive",
    )

    p = doc.add_paragraph(style="Lead")
    base.set_run_font(
        p.add_run("OpenClaw hữu ích cho điều phối agent, kênh chat và quy trình nội bộ; nó không phải model ASR và không làm giá OpenRouter rẻ hơn."),
        size=12,
        color=base.INK,
    )
    doc.add_page_break()


def assumptions_and_cost(doc: Document):
    base.add_heading(doc, "1. Phạm vi và giả định", 1)
    for text in [
        "10.000 người dùng × 10 phút = 100.000 phút audio trong một tháng.",
        "Tải được phân bổ đều trong 30 ngày để tận dụng quota miễn phí hằng ngày của Workers AI.",
        "Cả hai cách đều dùng Whisper Large V3 Turbo; chất lượng thực tế vẫn phải kiểm thử trên audio tiếng Việt vùng miền.",
        "Chi phí OpenClaw gồm model OpenRouter và một VPS nhỏ 5-20 USD/tháng. Đây là khoảng lập ngân sách; nếu dùng máy sẵn có, có thể giảm phần VPS nhưng tăng rủi ro vận hành.",
        "Không gồm thuế, lưu trữ audio, băng thông của hệ thống khác, nhân sự và chi phí đánh giá dữ liệu.",
    ]:
        base.add_list_item(doc, text)

    base.add_heading(doc, "2. So sánh chi phí chính", 1)
    rows = [
        ("Giá model", "0,00051 USD/phút", "0,04 USD/giờ = 0,0006667 USD/phút"),
        ("Chi phí model trước ưu đãi", money(CF_GROSS_AI), money(OR_MODEL_COST)),
        ("Quota miễn phí", f"≈ {CF_MONTHLY_FREE_MINUTES:,.0f} phút/tháng".replace(",", "."), "Không đưa quota miễn phí vào tính toán"),
        ("Chi phí model sau quota", money(CF_AI_AFTER_FREE), money(OR_MODEL_COST)),
        ("Nền tảng / hạ tầng", "Workers Paid: 5,00 USD/tháng", "OpenClaw: 0 USD tiền phần mềm; VPS: 5-20 USD/tháng"),
        ("Tổng dự kiến", f"{money(CF_TOTAL)} ({vnd(CF_TOTAL)})", f"{money(OR_TOTAL_LOW)}-{money(OR_TOTAL_HIGH)} ({vnd(OR_TOTAL_LOW)}-{vnd(OR_TOTAL_HIGH)})"),
    ]
    base.add_table(
        doc,
        ["Khoản mục", "Cách 1: Cloudflare", "Cách 2: OpenClaw + OpenRouter"],
        rows,
        [2100, 3100, 4160],
        font_size=8.5,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    base.add_source_note(doc, "Bảng 1. Chi phí cho 100.000 phút/tháng; làm tròn đến hai chữ số thập phân.")

    doc.add_paragraph()
    shape = doc.add_picture(str(CHART_PATH), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_picture_alt(shape, "Biểu đồ so sánh chi phí", "Ba thanh chi phí: Cloudflare 52,70 USD; OpenClaw và OpenRouter từ 71,67 đến 86,67 USD.")
    base.add_caption(doc, "Hình 1. So sánh tổng chi phí tháng theo giả định chính.")

    add_callout(
        doc,
        "Chênh lệch",
        f"OpenRouter đắt hơn Cloudflare khoảng 30,7% nếu chỉ so giá model. Sau khi cộng Workers Paid và VPS OpenClaw, phương án OpenClaw + OpenRouter cao hơn khoảng {money(OR_TOTAL_LOW - CF_TOTAL)}-{money(OR_TOTAL_HIGH - CF_TOTAL)} mỗi tháng.",
        kind="caution",
    )
    doc.add_page_break()


def calculation_details(doc: Document):
    base.add_heading(doc, "3. Cách tính chi tiết", 1)
    base.add_heading(doc, "3.1. Cloudflare Workers Paid + Whisper Turbo", 2)
    lines = [
        f"Tổng audio: {USERS:,} × {MINUTES_PER_USER} = {TOTAL_MINUTES:,} phút.".replace(",", "."),
        f"Chi phí ASR trước quota: {TOTAL_MINUTES:,} × 0,00051 = {money(CF_GROSS_AI)}.".replace(",", "."),
        f"Quota miễn phí có giá trị 0,11 USD/ngày, tương đương khoảng {CF_FREE_MINUTES_PER_DAY:.1f} phút/ngày.".replace(".", ","),
        f"Trong 30 ngày: 51,00 - 3,30 = {money(CF_AI_AFTER_FREE)} tiền AI.",
        f"Cộng Workers Paid 5,00 USD: tổng {money(CF_TOTAL)} = {vnd(CF_TOTAL)}.",
    ]
    for line in lines:
        base.add_list_item(doc, line, bullet=False)
    add_text(doc, "Nếu tài khoản đã trả Workers Paid cho dịch vụ khác, chi phí tăng thêm riêng cho ASR chỉ khoảng 47,70 USD/tháng.", bold_lead="Nếu tài khoản đã trả Workers Paid")

    base.add_heading(doc, "3.2. OpenClaw + OpenRouter", 2)
    lines = [
        "OpenRouter: 0,04 USD/giờ = 0,0006667 USD/phút.",
        f"100.000 phút ÷ 60 × 0,04 = {money(OR_MODEL_COST)} tiền model.",
        f"VPS chạy OpenClaw: khoảng {money(OPENCLAW_VPS_LOW)}-{money(OPENCLAW_VPS_HIGH)}/tháng.",
        f"Tổng: {money(OR_TOTAL_LOW)}-{money(OR_TOTAL_HIGH)} = {vnd(OR_TOTAL_LOW)}-{vnd(OR_TOTAL_HIGH)}.",
    ]
    for line in lines:
        base.add_list_item(doc, line, bullet=False)
    add_text(doc, "OpenClaw không thay thế Whisper. Luồng thực tế vẫn là ứng dụng gửi audio đến OpenClaw, sau đó OpenClaw gọi model Whisper qua OpenRouter.", bold_lead="OpenClaw không thay thế Whisper.")

    base.add_heading(doc, "3.3. Hai cách hiểu khác của '10 phút/người'", 2)
    daily_cf = 3_000_000 * CF_RATE_PER_MIN - CF_FREE_VALUE_PER_DAY * DAYS + CF_WORKERS_PAID
    daily_or = 3_000_000 * OR_RATE_PER_MIN
    rows = [
        ("Tổng một lần, dồn trong một ngày", "100.000", money(CF_ONE_DAY_TOTAL), f"{money(OR_TOTAL_LOW)}-{money(OR_TOTAL_HIGH)}"),
        ("Tổng một tháng, phân bổ 30 ngày", "100.000", money(CF_TOTAL), f"{money(OR_TOTAL_LOW)}-{money(OR_TOTAL_HIGH)}"),
        ("10 phút/người/ngày trong 30 ngày", "3.000.000", money(daily_cf), f"{money(daily_or + OPENCLAW_VPS_LOW)}-{money(daily_or + OPENCLAW_VPS_HIGH)}"),
    ]
    base.add_table(doc, ["Kịch bản", "Phút/tháng", "Cloudflare", "OpenClaw + OpenRouter"], rows, [3150, 1500, 2050, 2660], font_size=8.5)
    base.add_source_note(doc, "Bảng 2. Kịch bản chính của báo cáo là dòng thứ hai.")
    doc.add_page_break()


def architecture_and_delivery(doc: Document):
    base.add_heading(doc, "4. Kiến trúc và vận hành", 1)
    base.add_heading(doc, "Cách 1 - Cloudflare", 2)
    add_text(doc, "Ứng dụng → backend hiện tại → Cloudflare Workers AI → Whisper Large V3 Turbo → transcript.")
    for text in [
        "Ít thay đổi nhất vì dự án hiện đã có lớp gọi Cloudflare.",
        "Không cần tự vận hành gateway AI hoặc máy chủ riêng cho ASR.",
        "Quota miễn phí được tính hằng ngày; khi vượt quota cần Workers Paid.",
        "Phù hợp với luồng API sản phẩm có số lượng request lớn và yêu cầu vận hành đơn giản.",
    ]:
        base.add_list_item(doc, text)

    base.add_heading(doc, "Cách 2 - OpenClaw + OpenRouter", 2)
    add_text(doc, "Ứng dụng → backend hiện tại → OpenClaw Gateway → OpenRouter STT → Whisper Large V3 Turbo → transcript.")
    for text in [
        "Thêm một lớp gateway cần triển khai, cập nhật, bảo mật, giám sát và sao lưu cấu hình.",
        "OpenClaw phù hợp hơn cho agent, kênh chat, voice note và workflow; không tạo lợi thế chi phí ASR trong kịch bản này.",
        "OpenRouter yêu cầu gửi audio đến endpoint transcription; chi phí thật có thể đọc từ trường usage.cost của response.",
        "Audio dài nên được chia thành đoạn ngắn. Tài liệu OpenRouter khuyến nghị chia các bản ghi dài hơn khoảng một phút vì upstream có thể timeout sau 60 giây.",
    ]:
        base.add_list_item(doc, text)

    base.add_heading(doc, "5. So sánh ngoài chi phí", 1)
    rows = [
        ("Độ phức tạp", "Thấp", "Trung bình-cao; thêm gateway và VPS"),
        ("Thời gian ra production", "Nhanh nhất", "Lâu hơn do tích hợp, bảo mật và load test"),
        ("Điểm lỗi", "Backend + Cloudflare", "Backend + OpenClaw + VPS + OpenRouter + provider"),
        ("Khả năng điều phối agent/kênh", "Không phải mục tiêu chính", "Mạnh; phù hợp Telegram/Discord/voice note/workflow"),
        ("Fine-tune bằng 10.000 audio", "Không thực hiện trực tiếp trên model này", "Không; OpenRouter là dịch vụ inference, OpenClaw là gateway"),
        ("Tối ưu chi phí ở quy mô này", "Tốt hơn", "Kém hơn"),
    ]
    base.add_table(doc, ["Tiêu chí", "Cloudflare", "OpenClaw + OpenRouter"], rows, [2100, 2900, 4360], font_size=8.5)
    base.add_source_note(doc, "Bảng 3. Đánh giá kiến trúc cho luồng nhận dạng giọng nói batch của ứng dụng.")
    doc.add_page_break()


def rollout_and_recommendation(doc: Document):
    base.add_heading(doc, "6. Kế hoạch triển khai", 1)
    rows = [
        ("Cloudflare - xác nhận cấu hình", "0,5 ngày", "Kiểm tra model, tài khoản Paid, usage dashboard và biến môi trường production"),
        ("Cloudflare - đo lường", "1-2 ngày", "Thêm/kiểm tra log phút audio, latency, lỗi, fallback và chi phí"),
        ("Cloudflare - kiểm thử tải", "1-2 ngày", "Thử tải theo đoạn audio thực tế, kiểm tra timeout và retry"),
        ("Cloudflare - rollout", "1-3 ngày", "Canary 5% → 25% → 100%; theo dõi lỗi và chi phí"),
        ("OpenClaw + OpenRouter - POC", "3-5 ngày", "Cài Gateway, cấu hình key, STT model và adapter backend"),
        ("OpenClaw + OpenRouter - production hardening", "1-2 tuần", "Auth, rate limit, secret, retry, logging, alert và backup"),
        ("OpenClaw + OpenRouter - load test/rollout", "1 tuần", "Chunking, concurrency, timeout, canary và fallback"),
    ]
    base.add_table(doc, ["Hạng mục", "Thời gian", "Đầu ra"], rows, [3150, 1450, 4760], font_size=8.5)
    base.add_source_note(doc, "Bảng 4. Thời gian là ước tính cho một kỹ sư đã quen TypeScript/Next.js; người mới với OpenClaw cần dự phòng thêm 30-50%.")

    base.add_heading(doc, "7. Khuyến nghị", 1)
    add_callout(doc, "Quyết định đề xuất", "Tiếp tục Cloudflare Workers Paid + Whisper Large V3 Turbo cho production. Không đưa OpenClaw vào đường đi chính của audio chỉ với mục tiêu giảm chi phí.", kind="positive")
    for text in [
        "Giữ cấu hình Cloudflare hiện tại và bổ sung dashboard theo dõi số phút, chi phí, latency và tỷ lệ fallback.",
        "Tạo bộ test tiếng Việt có đủ Bắc - Trung - Nam, trẻ em/người lớn, thiết bị và mức nhiễu; so WER/CER trước khi thay provider.",
        "Chỉ thử OpenClaw nếu dự án thật sự cần bot đa kênh, agent workflow hoặc xử lý voice note ngoài luồng ứng dụng chính.",
        "Nếu mục tiêu dài hạn là giảm mạnh chi phí và tận dụng 10.000 audio để huấn luyện, hướng đúng là fine-tune PhoWhisper và self-host; đây là phương án thứ ba, không phải OpenClaw + OpenRouter.",
    ]:
        base.add_list_item(doc, text)

    base.add_heading(doc, "8. Checklist quyết định trong 7 ngày", 1)
    for text in [
        "Xác nhận tài khoản Cloudflare production đang dùng Workers Paid.",
        "Đọc Workers AI Usage để lấy số phút/ngày và chi phí thật.",
        "Chạy 300-500 audio đại diện vùng miền qua Cloudflare và OpenRouter; so transcript, latency và lỗi.",
        "Bật cảnh báo ngân sách và giới hạn request ở backend.",
        "Chốt Cloudflare làm primary, OpenAI/OpenRouter chỉ làm fallback nếu kết quả benchmark chứng minh cần thiết.",
    ]:
        base.add_list_item(doc, text)
    doc.add_page_break()


def sources(doc: Document):
    base.add_heading(doc, "9. Nguồn tham khảo", 1)
    add_source(doc, "Cloudflare - Whisper Large V3 Turbo, giá 0,00051 USD/phút", "https://developers.cloudflare.com/workers-ai/models/whisper-large-v3-turbo/")
    add_source(doc, "Cloudflare - Workers AI pricing và quota 10.000 Neurons/ngày", "https://developers.cloudflare.com/workers-ai/platform/pricing/")
    add_source(doc, "Cloudflare - Workers Paid tối thiểu 5 USD/tháng", "https://developers.cloudflare.com/workers/platform/pricing/")
    add_source(doc, "OpenRouter - Whisper Large V3 Turbo, giá 0,04 USD/giờ", "https://openrouter.ai/openai/whisper-large-v3-turbo")
    add_source(doc, "OpenRouter - Speech-to-Text API", "https://openrouter.ai/docs/guides/overview/multimodal/stt")
    add_source(doc, "OpenClaw - OpenRouter provider và cấu hình STT", "https://docs.openclaw.ai/openrouter")
    add_source(doc, "OpenClaw - Audio and voice notes", "https://docs.openclaw.ai/nodes/audio")

    base.add_heading(doc, "Ghi chú phương pháp", 2)
    for text in [
        "Giá và quota được kiểm tra ngày 27/07/2026 và có thể thay đổi; cần xác nhận lại trước khi ký mua hoặc rollout.",
        "Mức VPS 5-20 USD/tháng là giả định lập ngân sách cho gateway nhỏ, không phải giá niêm yết của OpenClaw hay OpenRouter.",
        "Không giả định hai nhà cung cấp cho transcript giống hệt nhau dù cùng tên model; cần benchmark bằng cùng file audio và cùng cấu hình ngôn ngữ.",
        "Chi phí VND dùng tỷ giá kế hoạch 26.000 VNĐ/USD và chưa gồm thuế.",
    ]:
        base.add_list_item(doc, text)


def audit(doc: Document):
    assert len(doc.paragraphs) >= 55
    assert len(doc.tables) >= 7
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        assert tbl_pr.find(qn("w:tblW")) is not None
        assert tbl_pr.find(qn("w:tblInd")) is not None
        assert table._tbl.tblGrid is not None
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    for token in ("turn23", "turn24", "TODO", "PLACEHOLDER"):
        assert token not in text


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    make_cost_chart(CHART_PATH)

    doc = Document()
    base.configure_styles(doc)
    doc._bullet_num_id = base.add_numbering_definition(doc, bullet=True)
    doc._decimal_num_id = base.add_numbering_definition(doc, bullet=False)
    set_section_chrome(doc)

    props = doc.core_properties
    props.title = "So sánh Cloudflare và OpenClaw + OpenRouter cho 10.000 người dùng"
    props.subject = "Chi phí, kiến trúc và kế hoạch triển khai ASR tiếng Việt"
    props.author = "Nhóm dự án"
    props.keywords = "Cloudflare, OpenClaw, OpenRouter, Whisper, ASR, chi phí"
    props.comments = "Cập nhật ngày 27/07/2026"

    cover(doc)
    assumptions_and_cost(doc)
    calculation_details(doc)
    architecture_and_delivery(doc)
    rollout_and_recommendation(doc)
    sources(doc)

    audit(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
