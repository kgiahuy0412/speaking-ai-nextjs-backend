import sys

from docx import Document


path = sys.argv[1] if len(sys.argv) > 1 else r"C:\Users\DELL\Downloads\Build AI.docx"
doc = Document(path)
print(f"PARAGRAPHS={len(doc.paragraphs)} TABLES={len(doc.tables)} SECTIONS={len(doc.sections)}")
for index, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if text:
        print(f"P{index:03d} [{paragraph.style.name}] {text}")

for table_index, table in enumerate(doc.tables):
    print(f"\nTABLE {table_index} rows={len(table.rows)} cols={len(table.columns)}")
    for row_index, row in enumerate(table.rows):
        values = " || ".join(cell.text.replace("\n", " / ") for cell in row.cells)
        print(f"R{row_index:02d}: {values}")
