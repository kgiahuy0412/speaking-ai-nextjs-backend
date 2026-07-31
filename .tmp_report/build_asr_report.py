from __future__ import annotations

import math
import os
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
WORK_DIR = ROOT / ".tmp_report"
OUT_PATH = OUT_DIR / "bao-cao-ai-giong-noi-v2-chi-tiet-model-va-chi-phi.docx"
ARCH_PATH = WORK_DIR / "kien-truc-hybrid.png"
COST_PATH = WORK_DIR / "diem-hoa-von-chi-phi.png"

USD_TO_VND = 26_000
CF_RATE_EXACT = 0.00051
CF_NEURONS_PER_MIN = 46.63
CF_FREE_NEURONS_PER_DAY = 10_000
CF_FREE_MIN_PER_DAY = CF_FREE_NEURONS_PER_DAY / CF_NEURONS_PER_MIN
RUNPOD_A5000_PER_HOUR = 0.27
RUNPOD_L4_PER_HOUR = 0.39
RUNPOD_A40_PER_HOUR = 0.44
RUNPOD_4090_PER_HOUR = 0.69
RUNPOD_A100_PER_HOUR = 1.39
RUNPOD_SERVERLESS_16GB_PER_HOUR = 0.58
RUNPOD_SERVERLESS_24GB_PER_HOUR = 0.69
HOURS_PER_MONTH = 720

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
INK = "0B2545"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D0D5DD"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "EAF6EE"
GREEN = "1F6B3A"
PALE_GOLD = "FFF6DD"
GOLD = "8A6500"
PALE_RED = "FCECEC"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "000000"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, name: str = "Calibri", size: float | None = None,
                 color: str | None = None, bold: bool | None = None,
                 italic: bool | None = None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=MID_GRAY, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = 120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text: str, *, bold=False, color=BLACK, size=9,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)


def add_hyperlink(paragraph, text: str, url: str, color=BLUE, underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single" if underline else "none")
    r_pr.append(u)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    r_pr.append(sz)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_field(paragraph, field: str):
    run = paragraph.add_run()
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def add_numbering_definition(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rpr.append(rfonts)
    lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_list_item(
    doc: Document,
    text: str,
    *,
    bullet=True,
    bold_lead: str | None = None,
    num_id: int | None = None,
):
    p = doc.add_paragraph(style="Normal")
    resolved_num_id = num_id if num_id is not None else (
        doc._bullet_num_id if bullet else doc._decimal_num_id
    )
    apply_num(p, resolved_num_id)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def set_picture_alt(inline_shape, title: str, description: str):
    inline_shape._inline.docPr.set("title", title)
    inline_shape._inline.docPr.set("descr", description)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None,
             align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED, italic=True)
    return p


def add_source_note(doc: Document, text: str):
    p = doc.add_paragraph(style="Table Citation")
    r = p.add_run(text)
    set_run_font(r, size=8.5, color=MUTED, italic=True)
    return p


def add_callout(doc: Document, label: str, text: str, *, kind="info"):
    fill, accent = {
        "info": (PALE_BLUE, NAVY),
        "positive": (PALE_GREEN, GREEN),
        "caution": (PALE_GOLD, GOLD),
        "risk": (PALE_RED, RED),
    }[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=accent, size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, size=10.5, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths: Sequence[int], *, font_size=8.5,
              alignments: Sequence[int] | None = None,
              header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        set_cell_shading(header.cells[idx], header_fill)
        align = alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(header.cells[idx], value, bold=True, color=NAVY,
                      size=font_size, align=align)
    for r_idx, row_values in enumerate(rows):
        row = table.add_row()
        for idx, value in enumerate(row_values):
            if r_idx % 2 == 1:
                set_cell_shading(row.cells[idx], "FAFBFC")
            align = alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[idx], str(value), size=font_size, align=align)
    set_table_geometry(table, widths)
    return table


def format_usd(value: float) -> str:
    if value < 10:
        return f"{value:,.2f} USD"
    return f"{value:,.0f} USD"


def format_vnd(value_usd: float) -> str:
    return f"{value_usd * USD_TO_VND:,.0f} VND".replace(",", ".")


def draw_centered(draw: ImageDraw.ImageDraw, box, text: str, font, fill, spacing=6):
    x0, y0, x1, y1 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, align="center", spacing=spacing)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text(((x0 + x1 - w) / 2, (y0 + y1 - h) / 2), text,
                        font=font, fill=fill, align="center", spacing=spacing)


def make_architecture_diagram(path: Path):
    W, H = 1700, 950
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    font_path = r"C:\Windows\Fonts\arial.ttf"
    bold_path = r"C:\Windows\Fonts\arialbd.ttf"
    f_title = ImageFont.truetype(bold_path, 38)
    f_box = ImageFont.truetype(bold_path, 24)
    f_small = ImageFont.truetype(font_path, 21)
    f_note = ImageFont.truetype(font_path, 19)

    draw.text((70, 42), "Kiến trúc đích: self-host ưu tiên, Cloudflare bảo đảm chất lượng",
              font=f_title, fill="#17365D")

    boxes = {
        "client": (70, 165, 300, 285),
        "api": (380, 165, 630, 285),
        "pre": (710, 165, 980, 285),
        "router": (1060, 165, 1320, 285),
        "result": (1400, 165, 1630, 285),
        "local": (1030, 385, 1290, 520),
        "cf": (1370, 385, 1630, 520),
        "store": (110, 690, 390, 820),
        "label": (470, 690, 750, 820),
        "train": (830, 690, 1110, 820),
        "registry": (1190, 690, 1470, 820),
    }
    labels = {
        "client": "Ứng dụng\nweb / mobile",
        "api": "Next.js API\nhiện tại",
        "pre": "Chuẩn hóa audio\n16 kHz + trim silence",
        "router": "ASR Router\nfeature flag",
        "result": "Transcript\n+ confidence",
        "local": "ASR tự host\nWhisper Turbo / PhoWhisper",
        "cf": "Cloudflare fallback\nlarge-v3-turbo",
        "store": "Kho dữ liệu mã hóa\nchỉ audio đã opt-in",
        "label": "Gán nhãn 2 vòng\n+ kiểm tra ngôn ngữ",
        "train": "Fine-tune / đánh giá\ntheo vùng miền",
        "registry": "Model registry\ncanary + rollback",
    }
    fills = {
        "client": "#EAF2F8", "api": "#EAF2F8", "pre": "#EAF2F8",
        "router": "#FFF6DD", "result": "#EAF6EE", "local": "#EAF6EE",
        "cf": "#FCECEC", "store": "#F2F4F7", "label": "#F2F4F7",
        "train": "#F2F4F7", "registry": "#F2F4F7",
    }
    outlines = {
        "router": "#8A6500", "local": "#1F6B3A", "cf": "#9B1C1C",
        "result": "#1F6B3A",
    }
    for key, box in boxes.items():
        draw.rounded_rectangle(box, radius=22, fill=fills[key],
                               outline=outlines.get(key, "#2E74B5"), width=4)
        draw_centered(draw, box, labels[key], f_box, "#17365D")

    def arrow(a, b, color="#667085", width=5):
        draw.line([a, b], fill=color, width=width)
        ang = math.atan2(b[1] - a[1], b[0] - a[0])
        size = 18
        for delta in (2.55, -2.55):
            p = (b[0] + size * math.cos(ang + delta), b[1] + size * math.sin(ang + delta))
            draw.line([b, p], fill=color, width=width)

    arrow((300, 225), (380, 225))
    arrow((630, 225), (710, 225))
    arrow((980, 225), (1060, 225))
    arrow((1320, 225), (1400, 225))
    arrow((1190, 285), (1160, 385), color="#1F6B3A")
    arrow((1290, 452), (1370, 452), color="#9B1C1C")
    arrow((1500, 385), (1515, 285), color="#9B1C1C")
    draw.text((1305, 417), "timeout /\nđộ tin cậy thấp", font=f_note, fill="#9B1C1C")

    draw.line([(850, 610), (850, 640)], fill="#98A2B3", width=3)
    draw.text((80, 610), "Vòng đời dữ liệu huấn luyện tách biệt với dữ liệu vận hành",
              font=f_small, fill="#667085")
    arrow((390, 755), (470, 755))
    arrow((750, 755), (830, 755))
    arrow((1110, 755), (1190, 755))
    arrow((1330, 690), (1165, 520), color="#2E74B5")
    draw.text((1380, 845), "Mỗi phiên bản đều có WER/CER, P95 và rollback",
              font=f_note, fill="#667085")
    img.save(path, quality=95)


def make_cost_chart(path: Path):
    W, H = 1700, 900
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    font_path = r"C:\Windows\Fonts\arial.ttf"
    bold_path = r"C:\Windows\Fonts\arialbd.ttf"
    f_title = ImageFont.truetype(bold_path, 38)
    f_axis = ImageFont.truetype(font_path, 21)
    f_label = ImageFont.truetype(bold_path, 23)
    left, top, right, bottom = 150, 130, 1550, 740
    draw.text((70, 40), "Điểm hòa vốn tham khảo - 10 giây/clip, 30 ngày/tháng",
              font=f_title, fill="#17365D")

    max_clips = 5_000_000
    max_cost = 480
    free_month = CF_FREE_MIN_PER_DAY * 30

    def px_x(clips):
        return left + (right - left) * clips / max_clips

    def px_y(cost):
        return bottom - (bottom - top) * cost / max_cost

    for cost in range(0, 481, 80):
        y = px_y(cost)
        draw.line([(left, y), (right, y)], fill="#E4E7EC", width=2)
        draw.text((70, y - 13), f"{cost} USD", font=f_axis, fill="#667085")
    for m in range(0, 6):
        x = px_x(m * 1_000_000)
        draw.line([(x, top), (x, bottom)], fill="#F2F4F7", width=2)
        draw.text((x - 30, bottom + 18), f"{m}M", font=f_axis, fill="#667085")
    draw.line([(left, top), (left, bottom), (right, bottom)], fill="#344054", width=4)

    points = []
    for i in range(101):
        clips = max_clips * i / 100
        audio_min = clips * 10 / 60
        if audio_min <= free_month:
            cost = 0
        else:
            cost = 5 + (audio_min - free_month) * CF_RATE_EXACT
        points.append((px_x(clips), px_y(cost)))
    draw.line(points, fill="#2E74B5", width=7)

    a5000 = RUNPOD_A5000_PER_HOUR * HOURS_PER_MONTH
    draw.line([(left, px_y(a5000)), (right, px_y(a5000))], fill="#1F6B3A", width=6)
    draw.rectangle((left, px_y(350), right, px_y(250)), fill="#FFF6DD", outline=None)
    draw.text((right - 560, px_y(350) + 18), "TCO tự host thực tế ước tính: 250-350 USD/tháng",
              font=f_axis, fill="#8A6500")
    # redraw primary lines on top of TCO band
    draw.line(points, fill="#2E74B5", width=7)
    draw.line([(left, px_y(a5000)), (right, px_y(a5000))], fill="#1F6B3A", width=6)

    cross_clips = ((a5000 - 5) / CF_RATE_EXACT + free_month) * 6
    cx = px_x(cross_clips)
    cy = px_y(a5000)
    draw.ellipse((cx - 10, cy - 10, cx + 10, cy + 10), fill="#9B1C1C")
    draw.line([(cx, cy), (cx + 110, cy - 90)], fill="#9B1C1C", width=3)
    draw.text((cx + 120, cy - 115), f"~{cross_clips/1_000_000:.2f}M clip/tháng\n(hòa vốn GPU compute-only)",
              font=f_label, fill="#9B1C1C")
    draw.text((left + 35, px_y(a5000) - 42), f"A5000 chạy 24/7: {a5000:.0f} USD/tháng (chưa CPU/storage/ops)",
              font=f_label, fill="#1F6B3A")
    draw.text((left + 35, px_y(60)), "Cloudflare: gói Paid + phút vượt miễn phí", font=f_label, fill="#2E74B5")
    draw.text((left, 790), "Ghi chú: đây là mô hình chi phí, không phải benchmark hiệu năng. Tự host chỉ rẻ khi GPU đủ tải; traffic ngắt quãng làm điểm hòa vốn xấu hơn.",
              font=f_axis, fill="#667085")
    img.save(path, quality=95)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)

    for name in ["Report Title", "Report Subtitle", "Kicker", "Lead", "Table Citation", "Small Note"]:
        if name not in styles:
            styles.add_style(name, 1)
    title = styles["Report Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.line_spacing = 1.0
    subtitle = styles["Report Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = rgb(DARK_BLUE)
    subtitle.paragraph_format.space_after = Pt(18)
    kicker = styles["Kicker"]
    kicker.font.name = "Calibri"
    kicker.font.size = Pt(10)
    kicker.font.bold = True
    kicker.font.color.rgb = rgb(GOLD)
    kicker.paragraph_format.space_after = Pt(8)
    lead = styles["Lead"]
    lead.font.name = "Calibri"
    lead.font.size = Pt(12)
    lead.font.color.rgb = rgb(INK)
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.line_spacing = 1.2
    citation = styles["Table Citation"]
    citation.font.name = "Calibri"
    citation.font.size = Pt(8.5)
    citation.font.italic = True
    citation.font.color.rgb = rgb(MUTED)
    citation.paragraph_format.space_before = Pt(4)
    citation.paragraph_format.space_after = Pt(4)
    small = styles["Small Note"]
    small.font.name = "Calibri"
    small.font.size = Pt(9)
    small.font.color.rgb = rgb(MUTED)
    small.paragraph_format.space_after = Pt(4)


def configure_sections(doc: Document):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        r = hp.add_run("BÁO CÁO PHƯƠNG ÁN ASR TIẾNG VIỆT  |  NỘI BỘ")
        set_run_font(r, size=8.5, color=MUTED, bold=True)

        first_header = section.first_page_header
        fp = first_header.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fr = fp.add_run("NGHIÊN CỨU KỸ THUẬT VÀ CHI PHÍ  |  27.07.2026")
        set_run_font(fr, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        r1 = p.add_run("Trang ")
        set_run_font(r1, size=9, color=MUTED)
        add_page_field(p, "PAGE")
        r2 = p.add_run(" / ")
        set_run_font(r2, size=9, color=MUTED)
        add_page_field(p, "NUMPAGES")

        first_footer = section.first_page_footer
        ffp = first_footer.paragraphs[0]
        ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = ffp.add_run("Tài liệu làm cơ sở ra quyết định - giá và giả định cần được cập nhật trước khi mua hạ tầng")
        set_run_font(rr, size=8.5, color=MUTED, italic=True)


def add_cover(doc: Document):
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph(style="Kicker")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("BÁO CÁO ĐỀ XUẤT PHƯƠNG ÁN")
    p = doc.add_paragraph(style="Report Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("AI nhận dạng giọng nói tiếng Việt\ncho ứng dụng luyện nói")
    p = doc.add_paragraph(style="Report Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Tối ưu chi phí Cloudflare mà vẫn giữ chất lượng ở mức tương đương\nBản V2: cách làm, triển khai và chi phí theo từng model")

    p = doc.add_paragraph(style="Lead")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Kết luận chính: chưa nên train từ đầu. Hãy đo đúng chi phí, xây bộ kiểm thử chuẩn, rồi chạy self-host ở chế độ shadow/hybrid trước khi quyết định fine-tune.")
    set_run_font(r, size=12, color=INK, bold=True)

    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
    meta = add_table(
        doc,
        ["Phạm vi", "Cơ sở đánh giá"],
        [
            ["ASR tiếng Việt, 10.000 audio, đa vùng miền", "Audit mã nguồn hiện tại + nguồn chính thức"],
            ["Đơn vị tiền", "USD và VND; tỷ giá kế hoạch 1 USD = 26.000 VND"],
            ["Ngày chốt thông tin", "27/07/2026"],
        ],
        [2500, 6860],
        font_size=9.5,
        header_fill=PALE_BLUE,
    )
    meta.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_executive_summary(doc: Document):
    add_heading(doc, "1. Kết luận điều hành", 1)
    p = doc.add_paragraph(style="Lead")
    p.add_run(
        "Với mức giá công bố hiện tại của Cloudflare, bài toán này trước hết là bài toán đo lường và chất lượng dữ liệu - chưa phải bài toán tự huấn luyện để tiết kiệm tiền."
    )
    add_callout(
        doc,
        "Khuyến nghị",
        "Giữ Cloudflare Whisper large-v3-turbo làm chuẩn và đường fallback trong 8-10 tuần đầu; xây golden set 1.500-2.000 clip, benchmark Whisper large-v3-turbo tự host, PhoWhisper-small và PhoWhisper-medium. Fine-tune PhoWhisper-small trước; chỉ nâng lên medium hoặc custom Turbo khi small không đạt quality gate.",
        kind="positive",
    )
    add_body(
        doc,
        "Cloudflare công bố đơn giá chính xác 0,00051 USD/phút cho model đang dùng. Vì vậy, 10.000 clip dài trung bình 10 giây chỉ tương đương khoảng 1.667 phút audio và 0,85 USD chi phí suy luận gộp trước hạn mức miễn phí. Nếu lượng dùng được phân bổ đều và toàn bộ 10.000 Neurons/ngày chưa bị các model khác dùng hết, phần ASR này có thể nằm trong hạn mức miễn phí. [1][2]",
    )
    add_body(
        doc,
        "Tự host một GPU A5000 giá tham khảo 0,27 USD/giờ sẽ tốn khoảng 194 USD/tháng nếu chạy liên tục, chưa gồm CPU, storage, egress, giám sát và nhân sự vận hành. Điểm hòa vốn compute-only tương ứng khoảng 2,27 triệu clip 10 giây/tháng; điểm hòa vốn TCO thực tế nên dùng vùng 3-5 triệu clip/tháng. [9]",
    )
    add_body(
        doc,
        "10.000 audio vẫn có giá trị lớn để tạo bộ đánh giá miền nghiệp vụ và fine-tune một model có sẵn. Tuy nhiên, số lượng file không đủ để mô tả dữ liệu: cần biết tổng giờ, số người nói, vùng miền, độ tuổi, thiết bị, nhiễu và chất lượng transcript. 10.000 clip × 10 giây chỉ khoảng 27,8 giờ - rất khác với 10.000 clip × 60 giây là 166,7 giờ.",
    )
    add_heading(doc, "Quyết định đề xuất", 2)
    decisions = [
        ("Trong 30 ngày", "Không mua GPU và không train. Thêm cost telemetry theo model, tạo consent flow, xây 1.500 clip golden set và đo Cloudflare baseline."),
        ("Trong 60 ngày", "Benchmark ba ứng viên: Cloudflare, Whisper Turbo tự host, PhoWhisper-medium; chạy shadow traffic, không ảnh hưởng người dùng."),
        ("Trong 90 ngày", "Nếu vượt quality gate, triển khai hybrid 5%-50% với Cloudflare fallback. Nếu không vượt, dừng self-host và tối ưu pipeline managed."),
        ("Chỉ fine-tune", "Khi lỗi tập trung rõ theo vùng miền/trẻ em/từ vựng miền và có dữ liệu opt-in đã gán nhãn tốt."),
    ]
    add_table(doc, ["Mốc", "Quyết định / kết quả cần có"], decisions, [1900, 7460], font_size=9)
    add_source_note(doc, "Quality gate chi tiết ở Mục 9; ngân sách ở Mục 8.")


def add_current_state(doc: Document):
    add_heading(doc, "2. Hiện trạng kỹ thuật của dự án", 1)
    add_body(
        doc,
        "Audit mã nguồn cho thấy hệ thống hiện tại đã có cấu trúc phù hợp để chuyển sang hybrid mà không cần viết lại toàn bộ pipeline.",
    )
    rows = [
        ("ASR chính", "@cf/openai/whisper-large-v3-turbo; language=vi; task=transcribe; VAD bật; condition_on_previous_text=false; prompt hướng tới câu ngắn của trẻ em."),
        ("Fallback", "Nếu Cloudflare lỗi/timeout/chất lượng bất thường, pipeline chuyển sang OpenAI ASR."),
        ("Sau ASR", "Transcript đi qua bước sửa lỗi/ngôn ngữ, sau đó LLM sinh câu tiếng Anh và TTS tạo audio. Vì vậy hóa đơn Cloudflare không chỉ có ASR."),
        ("Giới hạn upload", "Cloudflare audio mặc định 10 MB; audio session của ứng dụng mặc định 16 MB."),
        ("Lưu dữ liệu", "Chunk audio bị xóa ngay sau finalize thành công; kết quả finalize chỉ giữ theo TTL mặc định 86.400 giây. Hiện chưa có kho dữ liệu train bền vững."),
        ("Quan sát", "Có log latency/provider/fallback và báo cáo P50/P95, nhưng chưa có transcript chuẩn để đo WER/CER theo vùng miền."),
    ]
    add_table(doc, ["Hạng mục", "Phát hiện"], rows, [1900, 7460], font_size=8.7)
    add_source_note(
        doc,
        "Nguồn nội bộ: src/lib/ai/cloudflareWorkersAi.ts; src/lib/ai/asr.ts; src/lib/ai/pipeline.ts; src/lib/storage/audioSessions.ts; src/lib/storage/config.ts (audit ngày 27/07/2026).",
    )
    add_callout(
        doc,
        "Khoảng trống cần xử lý trước khi train",
        "Phải tách dữ liệu vận hành khỏi dữ liệu huấn luyện. Audio chỉ được đưa vào kho train khi người dùng/đại diện hợp pháp đã đồng ý rõ mục đích; cần retention, quyền rút lại, mã hóa, log truy cập và xóa theo yêu cầu.",
        kind="risk",
    )
    add_heading(doc, "Chi phí nào thực sự đang chiếm hóa đơn?", 2)
    add_body(
        doc,
        "Workers AI dùng một pool 10.000 Neurons miễn phí/ngày cho tổng các model. Dự án có thể gọi ASR, model text và TTS; một lỗi ASR còn có thể kích hoạt fallback. Do đó cần báo cáo cost theo provider + model + endpoint + phút audio/character/token, không chỉ nhìn tổng tiền tài khoản. [2]",
    )
    add_list_item(doc, "ASR: phút audio thực tế sau khi cắt im lặng; số request; tỷ lệ retry/fallback.")
    add_list_item(doc, "LLM: input/output tokens, cache hit của rule/text cache, model đang chạy.")
    add_list_item(doc, "TTS: số ký tự hoặc phút đầu ra, cache hit của audio đã sinh.")
    add_list_item(doc, "Hạ tầng ngoài AI: Worker requests/CPU, blob/object storage, database và egress.")
    add_body(
        doc,
        "Ví dụ, Cloudflare công bố Aura-1 ở mức 0,015 USD/1.000 ký tự. Nếu cấu hình TTS thực tế dùng Aura-1, 10.000 câu × 60 ký tự tương đương khoảng 9 USD - lớn hơn ví dụ ASR 10 giây/clip (0,85 USD). Điều này phải được xác nhận bằng dashboard và biến môi trường thực tế trước khi ưu tiên dự án ASR. [2]",
    )


def add_cost_analysis(doc: Document):
    add_heading(doc, "3. Mô hình chi phí và điểm hòa vốn", 1)
    add_callout(
        doc,
        "Cách đọc bảng",
        "Chi phí dưới đây là gross inference theo 0,00051 USD/phút để so sánh công bằng. Cloudflare còn miễn phí 10.000 Neurons/ngày (~214,45 phút ASR/ngày nếu pool chưa dùng cho model khác). Workers Paid có mức tối thiểu 5 USD/tháng khi cần vượt hạn mức. [1][2][5]",
        kind="info",
    )
    rows = []
    for avg_seconds in (5, 10, 15, 30, 60):
        minutes = 10_000 * avg_seconds / 60
        gross = minutes * CF_RATE_EXACT
        one_day_overage = max(0, minutes - CF_FREE_MIN_PER_DAY) * CF_RATE_EXACT
        rows.append((
            f"{avg_seconds} giây",
            f"{minutes:,.0f}".replace(",", "."),
            format_usd(gross),
            format_vnd(gross),
            f"5 USD + {one_day_overage:.2f} USD" if minutes > CF_FREE_MIN_PER_DAY else "0 USD",
        ))
    add_table(
        doc,
        ["Độ dài TB/clip", "Tổng phút", "Gross inference", "Quy đổi", "Nếu xử lý cả 10k trong 1 ngày"],
        rows,
        [1450, 1250, 1800, 1900, 2960],
        font_size=8.3,
        alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 5,
    )
    add_source_note(doc, "Bảng 1. Tính theo giá model page 0,00051 USD/phút; chưa gồm thuế. Kịch bản một ngày cộng 5 USD gói Paid và trừ 214,45 phút miễn phí.")

    add_heading(doc, "So sánh theo lưu lượng tháng (10 giây/clip)", 2)
    monthly_rows = []
    free_month = CF_FREE_MIN_PER_DAY * 30
    for clips in (10_000, 100_000, 1_000_000, 2_300_000, 5_000_000):
        minutes = clips / 6
        gross = minutes * CF_RATE_EXACT
        even_paid = 0 if minutes <= free_month else 5 + (minutes - free_month) * CF_RATE_EXACT
        monthly_rows.append((
            f"{clips:,.0f}".replace(",", "."),
            f"{minutes:,.0f}".replace(",", "."),
            format_usd(gross),
            format_usd(even_paid),
        ))
    add_table(
        doc,
        ["Clip/tháng", "Phút audio", "Gross inference", "Ước tính Paid nếu phân bổ đều"],
        monthly_rows,
        [1900, 1900, 2500, 3060],
        font_size=8.8,
        alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 4,
    )
    add_source_note(doc, "Bảng 2. Cột Paid giả định 30 ngày, pool free chưa bị model khác sử dụng. Nếu tài khoản đã trả gói Workers Paid vì dịch vụ khác, chỉ nên so phần chi phí tăng thêm.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cost_chart = p.add_run().add_picture(str(COST_PATH), width=Inches(6.45))
    set_picture_alt(
        cost_chart,
        "Biểu đồ điểm hòa vốn chi phí ASR",
        "So sánh chi phí Cloudflare theo số clip 10 giây mỗi tháng với chi phí GPU A5000 chạy liên tục.",
    )
    add_caption(doc, "Hình 1. Cloudflare chỉ giao nhau với GPU dedicated khi lưu lượng đạt hàng triệu clip ngắn/tháng; TCO thực tế cao hơn đường compute-only.")

    add_heading(doc, "Công thức ra quyết định", 2)
    add_body(doc, "Chi phí Cloudflare ASR ≈ tổng phút audio tính phí × 0,00051 USD/phút + phí Workers Paid (nếu cần).")
    add_body(doc, "Chi phí tự host ≈ GPU active/idle + CPU/RAM + storage + egress + monitoring + on-call + chi phí xây dựng/duy trì model.")
    add_body(doc, "Điểm hòa vốn compute-only với A5000 24/7 ≈ 194,4 / 0,00051 = 381.176 phút/tháng, tương đương khoảng 2,29 triệu clip 10 giây. Sau khi tính miễn phí + 5 USD, kết quả xấp xỉ 2,27 triệu clip. TCO thực tế nên đặt gate cao hơn: 3-5 triệu clip/tháng hoặc yêu cầu phi tài chính rõ ràng.")


def add_data_analysis(doc: Document):
    add_heading(doc, "4. 10.000 audio có đủ để train không?", 1)
    add_body(
        doc,
        "Đủ để fine-tune và đánh giá theo miền nghiệp vụ nếu dữ liệu được gán nhãn tốt; không đủ để train một mô hình ASR mạnh từ đầu. Whisper được huấn luyện trên quy mô rất lớn, còn PhoWhisper đạt độ bền tiếng Việt nhờ fine-tune trên 843,79 giờ và dữ liệu riêng từ 26.000 người thuộc 63 tỉnh/thành. [6][7][10]",
    )
    hours_rows = []
    for seconds in (5, 10, 15, 30, 60):
        hours = 10_000 * seconds / 3600
        hours_rows.append((f"{seconds} giây", f"{hours:.1f} giờ", f"{hours*0.70:.1f} giờ", f"{hours*0.15:.1f} giờ", f"{hours*0.15:.1f} giờ"))
    add_table(doc, ["TB/clip", "Tổng", "Train 70%", "Validation 15%", "Test 15%"], hours_rows,
              [1500, 1965, 1965, 1965, 1965], font_size=8.7,
              alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 5)
    add_source_note(doc, "Bảng 3. Split phải theo speaker, không random theo file; một người nói không xuất hiện ở cả train và test.")

    add_heading(doc, "Điều kiện tối thiểu để 10.000 clip có giá trị", 2)
    requirements = [
        "Có transcript chuẩn do con người xác nhận; transcript Cloudflare chỉ dùng để pre-fill, không được coi là ground truth tự động.",
        "Có speaker_id giả danh để split theo người; mục tiêu tối thiểu khoảng 300 người nói, ưu tiên đúng phân khúc người dùng thực tế.",
        "Có nhãn vùng miền (Bắc/Trung/Nam và nhóm địa phương quan trọng), độ tuổi, giới tính tự khai nếu được phép, thiết bị, mức nhiễu và bối cảnh.",
        "Có quota cân bằng. Ví dụ khởi điểm: Bắc 30%, Trung 30%, Nam 30%, khác/không rõ 10%; điều chỉnh theo thị trường thật.",
        "Giữ nguyên dấu tiếng Việt và từ địa phương; ban hành một quy tắc thống nhất cho số, tên riêng, từ đệm, lặp từ và câu chưa hoàn chỉnh.",
        "Tạo golden test 1.500-2.000 clip, khóa lại và không dùng để chọn checkpoint/hyperparameter.",
    ]
    for item in requirements:
        add_list_item(doc, item)

    add_heading(doc, "Nguồn dữ liệu bổ sung", 2)
    add_body(
        doc,
        "PhoWhisper là điểm khởi đầu tốt vì đã fine-tune trên 844 giờ tiếng Việt đa giọng. Mozilla Common Voice Vietnamese v23 có 18.777 clip, 22 giờ, 361 người nói nhưng chỉ 6,3 giờ validated và giấy phép CC0; phù hợp để tăng độ đa dạng, không thay thế dữ liệu hội thoại trẻ em của sản phẩm. VIVOS có khoảng 15 giờ nhưng giấy phép CC BY-NC-SA 4.0, có ràng buộc phi thương mại - không nên đưa vào sản phẩm thương mại nếu chưa được tư vấn pháp lý. [6][11][12]",
    )


def add_options(doc: Document):
    add_heading(doc, "5. Các phương án triển khai", 1)
    add_body(doc, "Các phương án dưới đây được xếp theo mức độ rủi ro tăng dần. Phương án tối ưu chi phí hiện tại là A; kiến trúc đích hợp lý nhất nếu quy mô tăng là B.")

    options = [
        ("A. Giữ Cloudflare, tối ưu pipeline", "2-4 tuần", "15-50 triệu VND", "0-5 USD/tháng + usage", "Thấp nhất", "Khuyến nghị ngay"),
        ("B. Hybrid: self-host + Cloudflare fallback", "6-10 tuần", "75-160 triệu VND", "20-350 USD/tháng + fallback", "Trung bình", "Khi cần kiểm soát/scale"),
        ("C. Fine-tune PhoWhisper/Whisper", "12-20 tuần", "180-450 triệu VND", "80-350 USD/tháng + MLOps", "Cao", "Khi benchmark chưa đạt"),
        ("D. On-device quantized", "10-16 tuần", "120-240 triệu VND", "Gần 0 phí ASR server", "Cao trên thiết bị yếu", "Android/native cụ thể"),
        ("E. Train từ đầu", "9-18+ tháng", ">2 tỷ VND", "Rất cao", "Rất cao", "Không đề xuất"),
    ]
    add_table(doc, ["Phương án", "Thời gian", "Chi phí dự án", "Vận hành", "Rủi ro", "Kết luận"], options,
              [2350, 1150, 1550, 1800, 1150, 1360], font_size=7.6,
              alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                          WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])
    add_source_note(doc, "Khoảng chi phí là ngân sách kế hoạch, gồm nhân sự theo đơn giá giả định ở Mục 8; chưa gồm thuế và chi phí tuyển dụng.")

    add_heading(doc, "A. Giữ Cloudflare và tối ưu trước khi train", 2)
    add_body(doc, "Mục tiêu là biết chính xác chi phí nào đáng tối ưu và cải thiện chất lượng bằng tham số/model prompt hiện có.")
    for item in [
        "Ghi duration_ms sau trim silence, model, neurons/ước tính, response status, retry, fallback, P50/P95; tạo dashboard theo ngày và endpoint.",
        "Không gửi đoạn im lặng dài. VAD đang bật ở Cloudflare, nhưng cắt im lặng trước upload mới giảm độ dài input và network; phải A/B test để không cắt mất âm đầu/cuối.",
        "Tối ưu beam_size và các threshold no-speech/log-prob/compression/hallucination trên golden set. Cloudflare hỗ trợ các tham số này; beam cao hơn có thể tăng chất lượng và latency. [1]",
        "Dùng hash idempotency để tránh retry trùng; cache chỉ khi audio thực sự trùng và chính sách riêng tư cho phép.",
        "Tách raw ASR khỏi transcript đã repair để đo model thật; đặc biệt theo dõi phủ định, đại từ/xưng hô, câu hỏi và từ địa phương.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "B. Hybrid self-host + Cloudflare fallback", 2)
    add_body(
        doc,
        "Chạy cùng model openai/whisper-large-v3-turbo giúp so sánh gần nhất với Cloudflare; model và code Whisper có giấy phép MIT. faster-whisper dùng CTranslate2, có thể nhanh hơn implementation Whisper gốc và hỗ trợ INT8, nhưng phải benchmark trên clip ngắn thực tế. [8][10]",
    )
    for item in [
        "Primary: Whisper large-v3-turbo tự host (để đạt parity trước), sau đó thử PhoWhisper-medium để giảm compute/tăng độ phù hợp tiếng Việt.",
        "Router fallback khi timeout, no-speech bất thường, log-prob thấp, script lạ, transcript rỗng hoặc rule kiểm tra nghĩa quan trọng thất bại.",
        "Shadow mode lưu cả output local và Cloudflare trên dữ liệu opt-in/đã ẩn danh; không thay đổi output cho người dùng trong giai đoạn đo.",
        "Canary 5% → 25% → 50% → 100%; rollback bằng feature flag trong vài phút.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "C. Fine-tune model có sẵn", 2)
    add_body(
        doc,
        "Ứng viên ưu tiên là PhoWhisper-medium hoặc Whisper large-v3-turbo với LoRA/PEFT. PhoWhisper-medium có 769M tham số và WER công bố gần PhoWhisper-large 1,55B trên nhiều benchmark, nên là điểm cân bằng hợp lý; số WER công bố không được so trực tiếp với Cloudflare nếu chưa chạy cùng test set. [6]",
    )
    for item in [
        "Bắt đầu LoRA/adapter để giảm VRAM và thời gian thử nghiệm; chỉ full fine-tune nếu adapter không đạt.",
        "Dùng 70/15/15 theo speaker; 3 seed; early stopping; lưu model card, dataset version, code commit và metric theo slice.",
        "Noise augmentation, speed perturbation nhẹ và room impulse response chỉ áp dụng khi giống môi trường thật; không augment vùng miền bằng biến đổi pitch giả tạo.",
        "Triển khai model đã chốt qua Transformers hoặc chuyển CTranslate2; kiểm tra sai lệch output sau conversion/quantization.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "D. On-device", 2)
    add_body(
        doc,
        "whisper.cpp hỗ trợ CPU, quantization, iOS, Android và WebAssembly. Phương án này có lợi cho offline/privacy và bỏ phí ASR server, nhưng kích thước model, nhiệt, pin và chênh lệch thiết bị làm rủi ro chất lượng/latency cao. Phù hợp hơn với Android native hoặc bộ câu lệnh đóng; chưa nên là đường chính cho web đa thiết bị. [13]",
    )


def add_architecture(doc: Document):
    add_heading(doc, "6. Kiến trúc đề xuất", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    architecture = p.add_run().add_picture(str(ARCH_PATH), width=Inches(6.45))
    set_picture_alt(
        architecture,
        "Kiến trúc ASR hybrid",
        "Luồng Next.js định tuyến sang ASR tự host và chuyển sang Cloudflare khi độ tin cậy thấp hoặc có lỗi; dữ liệu opt-in đi qua kho mã hóa, gán nhãn và MLOps.",
    )
    add_caption(doc, "Hình 2. Kiến trúc hybrid giữ được đường lui về Cloudflare trong khi tích lũy bằng chứng về chi phí và chất lượng.")
    add_heading(doc, "Thành phần cần bổ sung", 2)
    components = [
        ("Audio preprocessor", "Chuẩn hóa mono 16 kHz; trim silence an toàn; checksum; đo duration; không lưu nếu chưa opt-in."),
        ("ASR router", "Feature flag, timeout budget, quality heuristic, circuit breaker, retry idempotent, Cloudflare fallback."),
        ("ASR service", "FastAPI/gRPC; model warm; micro-batching; GPU queue; /health, /ready, /metrics; request size limit."),
        ("Quality service", "WER/CER định kỳ; critical-token error; phát hiện script lạ/no-speech/hallucination; slice dashboard."),
        ("Data plane", "Object storage mã hóa; manifest versioned; consent/retention/delete; annotation UI; audit log."),
        ("MLOps", "Dataset registry, model registry, experiment tracking, signed image, canary, rollback."),
    ]
    add_table(doc, ["Thành phần", "Trách nhiệm"], components, [2100, 7260], font_size=8.7)

    add_heading(doc, "Giao diện API tương thích", 2)
    add_body(doc, "Giữ response tương thích với code hiện tại: text, word_count, segments, vtt, model; bổ sung provider, confidence/quality_flags, latency_ms, fallback_used và model_version. Cách này cho phép thay adapter Cloudflare bằng adapter self-host mà không làm thay đổi pipeline LLM/TTS.")


def add_training_guide(doc: Document):
    add_heading(doc, "7. Quy trình training dành cho người mới", 1)
    training_num_id = add_numbering_definition(doc, bullet=False)
    steps = [
        ("Chốt mục tiêu", "Không dùng câu “hiểu 10.000 audio” làm metric. Chốt WER/CER, lỗi từ quan trọng, P95, tỷ lệ fallback và chi phí/1.000 phút."),
        ("Thu thập hợp lệ", "Consent riêng cho cải thiện/huấn luyện; speaker_id giả danh; retention; quyền rút lại; tách train storage khỏi production."),
        ("Làm sạch audio", "Decode được; 16 kHz mono; không clipping; loại file rỗng/trùng; đo SNR/độ dài; giữ bản gốc có kiểm soát."),
        ("Gán nhãn", "Cloudflare pre-fill → annotator nghe/sửa → reviewer kiểm tra 10-20%; adjudication cho bất đồng; version guideline."),
        ("Chia dữ liệu", "Split theo speaker và stratify vùng/tuổi/noise; khóa golden test; tránh leakage từ câu lặp của cùng một người."),
        ("Benchmark zero-shot", "Chạy Cloudflare, Whisper Turbo self-host, PhoWhisper-small/medium; cùng normalization, cùng test, cùng hardware profile."),
        ("Fine-tune", "LoRA trước; sweep nhỏ learning rate/batch; 3 seed; early stopping; ghi experiment; không nhìn test set."),
        ("Error analysis", "Lấy top lỗi theo vùng, tuổi, thiết bị, noise và từ quan trọng; bổ sung dữ liệu đúng lỗi, không chỉ train thêm epoch."),
        ("Tối ưu inference", "CTranslate2/INT8 nếu quality loss đạt gate; benchmark clip 3-15 giây và concurrency thật; đo cold start."),
        ("Triển khai an toàn", "Shadow → canary → hybrid; Cloudflare fallback; monitor drift; rollback; tái đánh giá hàng tháng."),
    ]
    for title, detail in steps:
        add_list_item(
            doc,
            f"{title}: {detail}",
            bullet=False,
            bold_lead=f"{title}: ",
            num_id=training_num_id,
        )

    add_heading(doc, "Quy tắc transcript tiếng Việt", 2)
    add_body(doc, "Trước khi thuê gán nhãn, cần chốt một trang guideline với ví dụ thật. Các quyết định phải nhất quán:")
    for item in [
        "Giữ nguyên dấu; NFC Unicode; không tự sửa ngữ pháp hoặc đổi từ địa phương sang từ phổ thông.",
        "Ghi từ đệm/lặp nếu nghe rõ; đánh dấu [không nghe rõ] theo chuẩn, không đoán.",
        "Chốt cách ghi số, chữ cái, tên riêng, tiếng Anh xen kẽ và âm thanh không lời.",
        "Ghi đúng phủ định, xưng hô, chủ thể, địa điểm, câu hỏi - đây là token quan trọng của sản phẩm.",
        "Tách raw transcript khỏi normalized transcript để WER không bị “đẹp giả” bởi bước hậu xử lý.",
    ]:
        add_list_item(doc, item)


def add_model_playbooks(doc: Document):
    add_heading(doc, "8. Cách làm, triển khai và chi phí theo từng model", 1)
    add_callout(
        doc,
        "Khuyến nghị triển khai",
        "Dùng Cloudflare hoặc Whisper Large V3 Turbo local làm teacher để tạo transcript nháp; fine-tune PhoWhisper-small trước; chỉ nâng lên PhoWhisper-medium khi small không đạt quality gate; giữ Cloudflare làm fallback. Không chạy đồng thời ba model self-host trong production.",
        kind="positive",
    )
    add_body(
        doc,
        "Các con số dưới đây là planning estimate cho 10.000 clip trung bình 10 giây (1.666,7 phút / 27,8 giờ audio), chia 70/15/15 theo người nói, 3 lượt thí nghiệm và tỷ giá 26.000 VND/USD. Giá GPU công bố ngày 17/07/2026; chưa gồm thuế, storage, egress và thời gian chờ do hết GPU. [1][2][8][9][10][16][17][18][19]",
    )
    model_rows = [
        ("Whisper Large V3 Turbo", "Teacher/baseline", "~809M; 1,62 GB; MIT", "24-48 GB khi LoRA", "Gần nhất với model Cloudflare hiện tại"),
        ("PhoWhisper-small", "Student tiết kiệm", "~244M; 971 MB; BSD-3", "12-24 GB khi LoRA", "Ứng viên production đầu tiên"),
        ("PhoWhisper-medium", "Student chất lượng", "~769M; 3,06 GB; BSD-3", "24 GB LoRA; 48-80 GB full", "Dùng khi small chưa đạt"),
        ("Cloudflare Whisper Turbo", "Managed fallback", "Không tải/training custom", "Không quản GPU", "Fallback, baseline và pre-label"),
    ]
    add_table(
        doc,
        ["Model", "Vai trò", "Quy mô / license", "GPU training", "Cách dùng"],
        model_rows,
        [2000, 1600, 1900, 1700, 2160],
        font_size=7.9,
    )
    add_source_note(doc, "Kích thước file từ model repository; số tham số theo cấu hình Whisper tương ứng. VRAM là vùng lập kế hoạch và phải được xác nhận bằng batch test thực tế. [10][16][17]")

    add_heading(doc, "8.1 openai/whisper-large-v3-turbo - teacher và baseline", 2)
    add_body(
        doc,
        "Đây là lựa chọn gần nhất với @cf/openai/whisper-large-v3-turbo mà dự án đang gọi. Bản Hugging Face có giấy phép MIT, khoảng 809M tham số và file trọng số 1,62 GB. Cloudflare-hosted chỉ cho inference; muốn fine-tune bằng 10.000 audio thì phải tải model mở và tự triển khai checkpoint của mình. [1][10]",
    )
    add_heading(doc, "Cách làm", 3)
    for item in [
        "Dùng Cloudflare hoặc model local để tạo transcript nháp cho 10.000 clip; người gán nhãn phải nghe và sửa trước khi dữ liệu được xem là ground truth.",
        "Giữ model gốc làm baseline bất biến. Nếu muốn thử fine-tune Turbo, tạo nhánh LoRA riêng với target q_proj/v_proj, rank 16-32, gradient checkpointing, FP16/BF16 và batch nhỏ.",
        "Không dùng output teacher chưa kiểm tra để train student; ưu tiên sửa phủ định, xưng hô, số, tên, địa điểm và từ khóa bài học.",
        "Chạy cùng golden test với Cloudflare, PhoWhisper-small và PhoWhisper-medium; giữ beam size, VAD và normalization giống nhau.",
    ]:
        add_list_item(doc, item)
    add_heading(doc, "Triển khai", 3)
    for item in [
        "Giai đoạn pre-label: chạy batch job theo giờ, scale về 0 sau khi xong; không cần giữ GPU 24/7.",
        "Nếu chọn Turbo làm production primary: merge adapter, chuyển checkpoint sang CTranslate2, thử INT8-FP16, đóng gói FastAPI/gRPC và giữ Cloudflare fallback.",
        "Dành 24 GB GPU cho LoRA rất tiết kiệm batch; dùng 48 GB A40/A6000 để giảm rủi ro OOM. Full fine-tune nên dùng 80 GB và chỉ làm khi LoRA thất bại.",
    ]:
        add_list_item(doc, item)
    turbo_costs = [
        ("License model", "MIT; không có phí license model", "0 USD", "0 VND"),
        ("Pre-label bằng Cloudflare", "10.000 clip x 10 giây; gross trước free allocation", "~0,85 USD", "~22.100 VND"),
        ("Pre-label local", "A5000 2-8 GPU-giờ, batch inference", "0,54-2,16 USD", "~14.000-56.000 VND"),
        ("LoRA tùy chọn", "A40 48 GB, 80-200 GPU-giờ, gồm nhiều run", "35-88 USD", "~0,92-2,29 triệu"),
        ("Full fine-tune", "A100 80 GB, 100-250 GPU-giờ", "139-348 USD", "~3,61-9,05 triệu"),
        ("Inference serverless", "24 GB, 50-200 active giờ/tháng", "34,5-138 USD/tháng", "~0,90-3,59 triệu/tháng"),
        ("Inference dedicated", "A5000 pod bật 24/7", "194,4 USD/tháng", "~5,05 triệu/tháng"),
    ]
    add_table(doc, ["Khoản", "Giả định", "USD", "VND"], turbo_costs, [2100, 3700, 1600, 1960], font_size=8.0)
    add_source_note(doc, "Chi phí compute chỉ là tiền GPU. Turbo nên là teacher/baseline; fine-tune nó chỉ đáng làm khi PhoWhisper không đạt hoặc cần giữ tối đa tính tương thích với Cloudflare.")

    add_heading(doc, "8.2 vinai/PhoWhisper-small - student tiết kiệm", 2)
    add_body(
        doc,
        "PhoWhisper-small dựa trên Whisper-small, được VinAI thích nghi cho tiếng Việt; repository khoảng 971 MB và giấy phép BSD-3-Clause. Đây là model nên fine-tune đầu tiên vì vòng lặp thử nghiệm nhanh, checkpoint nhỏ và có thể chạy CPU INT8 khi traffic thấp. [6][16]",
    )
    add_heading(doc, "Cách làm", 3)
    for item in [
        "Khởi động bằng zero-shot trên golden test. Chỉ train khi đã biết model sai ở vùng, nhóm tuổi hoặc loại câu nào.",
        "Fine-tune LoRA 5-10 epoch, learning rate ban đầu 1e-4, rank 16, alpha 32, dropout 0,05; early stopping theo WER validation.",
        "Chạy 3 seed; giữ checkpoint tốt nhất theo WER tổng và critical-token error, không theo training loss đơn thuần.",
        "Nếu small đạt quality gate, dừng ở đây; không chuyển lên medium chỉ vì model lớn hơn.",
    ]:
        add_list_item(doc, item)
    add_heading(doc, "Triển khai", 3)
    for item in [
        "POC: Transformers FP16 trên GPU 12-24 GB. Production: merge LoRA rồi chuyển CTranslate2 INT8-FP16.",
        "Traffic thấp: CPU INT8 4-8 vCPU có thể được benchmark; traffic tương tác/concurrency cao: serverless GPU 16 GB hoặc A5000.",
        "Đặt timeout ngắn hơn Cloudflare fallback; log model_version, latency, language probability, no-speech và fallback_used.",
    ]:
        add_list_item(doc, item)
    small_costs = [
        ("License model", "BSD-3-Clause", "0 USD", "0 VND"),
        ("LoRA khuyến nghị", "A5000 24 GB, 30-80 GPU-giờ, 3 run", "8,1-21,6 USD", "~0,21-0,56 triệu"),
        ("Full fine-tune tùy chọn", "A40 48 GB, 60-120 GPU-giờ", "26-53 USD", "~0,69-1,37 triệu"),
        ("CPU low traffic", "Planning estimate, cần benchmark máy thật", "30-100 USD/tháng", "~0,78-2,60 triệu/tháng"),
        ("Serverless GPU", "16 GB, 50-200 active giờ/tháng", "29-116 USD/tháng", "~0,75-3,02 triệu/tháng"),
        ("Dedicated GPU", "A5000 pod 24/7", "194,4 USD/tháng", "~5,05 triệu/tháng"),
        ("Pilot trọn gói", "Label + ML/backend + QA + deploy", "Planning budget", "~75-140 triệu VND"),
    ]
    add_table(doc, ["Khoản", "Giả định", "USD", "VND"], small_costs, [2100, 3700, 1600, 1960], font_size=8.0)
    add_callout(doc, "Điều kiện chọn small", "Chọn PhoWhisper-small làm primary nếu đạt quality gate, P95 dưới 2,5 giây ở concurrency thật và fallback ổn định dưới 5-10%. Đây là lựa chọn cân bằng chi phí tốt nhất.", kind="info")

    add_heading(doc, "8.3 vinai/PhoWhisper-medium - student chất lượng cao", 2)
    add_body(
        doc,
        "PhoWhisper-medium dựa trên Whisper-medium, repository khoảng 3,06 GB, giấy phép BSD-3-Clause và nhiều lớp hơn small. Model này có dư địa chất lượng cao hơn nhưng training, cold start và concurrency đắt hơn; chỉ mở nhánh này khi error analysis chỉ ra small thiếu năng lực model. [6][17]",
    )
    add_heading(doc, "Cách làm", 3)
    for item in [
        "Khởi tạo từ PhoWhisper-medium, dùng đúng dataset/split/normalization của small để so sánh công bằng.",
        "LoRA rank 16-32, alpha 32-64, batch 1-4, gradient accumulation, gradient checkpointing và mixed precision; theo dõi OOM từ bước đầu.",
        "Không train nhiều epoch mặc định: bắt đầu 3-5 epoch, mở rộng khi validation còn cải thiện và không tăng lỗi vùng miền.",
        "Chỉ full fine-tune trên 48-80 GB khi LoRA đã được tối ưu nhưng vẫn thiếu chất lượng có ý nghĩa sản phẩm.",
    ]:
        add_list_item(doc, item)
    add_heading(doc, "Triển khai", 3)
    for item in [
        "Production nên dùng GPU 24 GB, CTranslate2 và INT8-FP16; benchmark batch 1 cho latency và batch 4-16 cho throughput.",
        "Serverless phù hợp traffic biến động; dedicated A5000/L4 phù hợp khi tổng active compute vượt khoảng 282 giờ/tháng và cần tránh cold start.",
        "Giới hạn concurrency/queue để giữ VRAM headroom; circuit breaker chuyển ngay sang Cloudflare khi queue, timeout hoặc OOM.",
    ]:
        add_list_item(doc, item)
    medium_costs = [
        ("License model", "BSD-3-Clause", "0 USD", "0 VND"),
        ("LoRA A5000", "24 GB, 60-160 GPU-giờ, 3 run", "16-43 USD", "~0,42-1,12 triệu"),
        ("LoRA L4", "24 GB, 60-160 GPU-giờ", "23-62 USD", "~0,61-1,62 triệu"),
        ("Full fine-tune", "A100 80 GB, 80-200 GPU-giờ", "111-278 USD", "~2,89-7,23 triệu"),
        ("Serverless GPU", "24 GB, 50-200 active giờ/tháng", "34,5-138 USD/tháng", "~0,90-3,59 triệu/tháng"),
        ("Dedicated GPU", "A5000-L4 bật 24/7", "194-281 USD/tháng", "~5,05-7,30 triệu/tháng"),
        ("Pilot trọn gói", "Label + ML/backend + QA + deploy", "Planning budget", "~90-180 triệu VND"),
    ]
    add_table(doc, ["Khoản", "Giả định", "USD", "VND"], medium_costs, [2100, 3700, 1600, 1960], font_size=8.0)
    add_callout(doc, "Điều kiện nâng lên medium", "Chỉ nâng cấp nếu medium giảm lỗi có ý nghĩa sản phẩm so với small và phần tăng chi phí latency/compute được chấp nhận. Chênh lệch WER nhỏ nhưng không giảm critical-token error chưa đủ để đổi model.", kind="caution")

    add_heading(doc, "8.4 Cloudflare Whisper Turbo - managed baseline và fallback", 2)
    add_body(
        doc,
        "Cloudflare-hosted không phải model để đưa 10.000 audio vào fine-tune. Nó là dịch vụ inference: dùng làm baseline, teacher pre-fill và fallback. Model page công bố 0,00051 USD/phút; bảng platform làm tròn 0,0005 USD/phút và 46,63 neurons/phút, với 10.000 neurons miễn phí mỗi ngày trong pool chung. [1][2]",
    )
    add_heading(doc, "Cách triển khai trong dự án hiện tại", 3)
    for item in [
        "Giữ adapter Cloudflare hiện có; thêm adapter self-host và ASR router trong src/lib/ai/asr.ts.",
        "Router gọi local trước; fallback khi timeout, lỗi 5xx/OOM, không có speech, log probability thấp hoặc quality heuristic không đạt.",
        "Dùng idempotency key để tránh tính phí/retry kép; circuit breaker chuyển 100% sang Cloudflare khi local service không ổn định.",
        "Gửi cùng language=vi, task=transcribe, VAD và decoding settings tương đương để benchmark công bằng.",
    ]:
        add_list_item(doc, item)
    cf_rows = []
    for clips in (10_000, 100_000, 1_000_000, 5_000_000):
        minutes = clips / 6
        gross = minutes * CF_RATE_EXACT
        fallback = gross * 0.05
        cf_rows.append((
            f"{clips:,.0f}".replace(",", "."),
            f"{minutes:,.0f}".replace(",", "."),
            format_usd(gross),
            format_usd(fallback),
            format_vnd(fallback),
        ))
    add_table(
        doc,
        ["Clip 10 giây/tháng", "Phút audio", "100% Cloudflare", "Fallback 5%", "Fallback 5% (VND)"],
        cf_rows,
        [1900, 1700, 1900, 1700, 2160],
        font_size=8.2,
        alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 5,
    )
    add_source_note(doc, "Gross inference trước free allocation và phí Workers Paid tối thiểu 5 USD/tháng nếu tài khoản cần nâng cấp. Nếu tài khoản đã dùng Paid vì dịch vụ khác, so sánh theo chi phí tăng thêm. [1][2][5]")

    add_heading(doc, "8.5 Chi phí chung và thứ tự ra quyết định", 2)
    common_costs = [
        ("Transcript + QC", "125-250 giờ; 100.000-200.000 VND/giờ", "12,5-50 triệu VND"),
        ("Data/consent pipeline", "Lưu opt-in, mã hóa, manifest, retention/delete", "10-30 triệu VND"),
        ("ML benchmark/training", "Golden set, 3 model, error analysis, model card", "20-60 triệu VND"),
        ("Backend/MLOps", "ASR service, router, metrics, canary, rollback", "25-75 triệu VND"),
        ("QA/privacy/security", "Slice review, load test, DPIA/legal review", "10-30 triệu VND"),
    ]
    add_table(doc, ["Khoản chung", "Phạm vi", "Ngân sách"], common_costs, [2400, 4400, 2560], font_size=8.4)
    add_body(doc, "Không cộng chi phí chung nhiều lần khi so model. Ngân sách quản trị nên dùng: Cloudflare-only 16-53 triệu; hybrid PhoWhisper-small 75-140 triệu; hybrid PhoWhisper-medium 90-180 triệu; custom Turbo hoặc full fine-tune 180-450 triệu tùy phạm vi.")
    decision_num_id = add_numbering_definition(doc, bullet=False)
    decisions = [
        ("Tuần 1-4", "Dùng Cloudflare/Turbo tạo nhãn nháp, khóa golden test và chạy zero-shot cả ba model."),
        ("Tuần 5", "Fine-tune PhoWhisper-small bằng LoRA; dừng nếu đạt gate."),
        ("Tuần 6", "Chỉ fine-tune PhoWhisper-medium nếu small chưa đạt và lỗi có khả năng cải thiện bằng model lớn."),
        ("Tuần 7-10", "Chuyển model thắng sang CTranslate2, shadow/canary; Cloudflare giữ fallback."),
        ("Sau 30 ngày production", "So TCO thực tế; chỉ chuyển dedicated GPU nếu active compute đủ lớn và cold start ảnh hưởng UX."),
    ]
    for title, detail in decisions:
        add_list_item(doc, f"{title}: {detail}", bullet=False, bold_lead=f"{title}: ", num_id=decision_num_id)


def add_budget(doc: Document):
    add_heading(doc, "9. Ngân sách dự kiến", 1)
    add_callout(
        doc,
        "Giả định lập ngân sách",
        "Đơn giá nhân sự minh họa 1,0-3,0 triệu VND/person-day; gán nhãn 0,10-0,20 triệu VND/giờ. Hãy thay bằng rate nội bộ/nhà cung cấp. Tỷ giá kế hoạch 26.000 VND/USD.",
        kind="caution",
    )
    budget_rows = [
        ("A. Audit + tối ưu Cloudflare", "10-20", "15-50", "0,5-3", "16-53 triệu VND"),
        ("B. Hybrid POC → canary", "35-60", "52-180", "3-15", "75-160 triệu VND (mục tiêu)"),
        ("C. Fine-tune + production", "70-120", "105-360", "GPU 3-26; label 12-50; privacy/QA 10-30", "180-450 triệu VND"),
        ("D. On-device", "50-80", "75-240", "thiết bị/QA 15-40", "120-240 triệu VND"),
    ]
    add_table(doc, ["Gói", "Person-days", "Nhân sự (triệu VND)", "Khác (triệu VND)", "Ngân sách quản trị"], budget_rows,
              [2200, 1150, 1900, 2100, 2010], font_size=8.1,
              alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                          WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])
    add_source_note(doc, "Ngân sách quản trị có thể không bằng tổng min/max từng cột vì phản ánh phạm vi mục tiêu và dự phòng; cần lập SOW chi tiết trước khi ký.")

    add_heading(doc, "Chi phí compute tham khảo", 2)
    compute_rows = [
        ("RunPod RTX A5000 pod", "24 GB", "0,27 USD/giờ", "~194 USD/tháng nếu 24/7", "POC/inference tiết kiệm"),
        ("RunPod L4 pod", "24 GB", "0,39 USD/giờ", "~281 USD/tháng nếu 24/7", "Inference data-center GPU"),
        ("RunPod RTX 4090 pod", "24 GB", "0,69 USD/giờ", "~497 USD/tháng nếu 24/7", "Experiment/fine-tune nhanh"),
        ("RunPod A100 80 GB", "80 GB", "1,39 USD/giờ", "Bật theo giờ", "Full fine-tune/model lớn"),
    ]
    add_table(doc, ["GPU", "VRAM", "Giá công bố", "Nếu luôn bật", "Vai trò"], compute_rows,
              [2100, 900, 1700, 2200, 2460], font_size=8.5,
              alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                          WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    add_source_note(doc, "Nguồn giá: RunPod pricing, cập nhật 17/07/2026. Giá/availability thay đổi theo loại cloud và vùng; chưa gồm storage/egress. [9]")

    add_heading(doc, "Chi phí gán nhãn 10.000 clip", 2)
    add_body(doc, "Với clip ngắn, giả định thao tác nghe-sửa-QA 45-90 giây/clip: 125-250 giờ lao động. Ở đơn giá 100.000-200.000 VND/giờ, chi phí 12,5-50 triệu VND. Có thể giảm thời gian bằng transcript pre-fill, nhưng reviewer vẫn phải nghe audio; cần cộng 10-20% adjudication và quản lý chất lượng.")


def add_timeline(doc: Document):
    add_heading(doc, "10. Kế hoạch triển khai chi tiết", 1)
    add_body(doc, "Lộ trình có các cổng dừng để tránh tiếp tục đầu tư khi bằng chứng không ủng hộ. Nhánh A kết thúc ở tuần 4; nhánh B ở tuần 10; nhánh C kéo dài đến tuần 20.")
    roadmap = [
        ("Tuần 1", "Đo baseline", "Dashboard minutes/neurons/model; p50/p95; retry/fallback; bill reconciliation", "Cost baseline đã đối soát"),
        ("Tuần 2-3", "Dữ liệu & pháp lý", "Consent flow; data manifest; guideline; gán nhãn 1.500 golden clip", "Golden set v1 + DPIA/retention"),
        ("Tuần 4", "Benchmark model", "Cloudflare vs Whisper Turbo vs PhoWhisper-small/medium; WER/CER/slice/latency/cost", "Gate 1: dừng hay làm hybrid"),
        ("Tuần 5-6", "ASR service", "Container; preprocessor; API compatible; metrics; auth; queue; load test", "Endpoint staging ổn định"),
        ("Tuần 7-8", "Shadow", "100% request chạy local song song theo mẫu opt-in; Cloudflare vẫn trả kết quả", "Báo cáo parity + capacity"),
        ("Tuần 9", "Canary 5%-25%", "Router confidence; circuit breaker; rollback; support playbook", "Gate 2: mở rộng hay rollback"),
        ("Tuần 10", "Canary 50%", "Theo dõi quality/latency/cost; chốt TCO", "Quyết định production hybrid"),
        ("Tuần 5-8*", "Label phần còn lại", "Gán nhãn 8.500 clip; QC; re-balance vùng/tuổi/noise", "Dataset train v1"),
        ("Tuần 9-12*", "Fine-tune", "LoRA/PEFT; sweep nhỏ; 3 seed; error analysis; model card", "Candidate model v1"),
        ("Tuần 13-16*", "Tối ưu & hardening", "CTranslate2/INT8; security; autoscale; failover; regression suite", "Release candidate"),
        ("Tuần 17-20*", "Rollout", "Shadow → 5 → 25 → 50%; audit drift và incident", "Gate 3: 100% hay hybrid lâu dài"),
    ]
    add_table(doc, ["Thời gian", "Workstream", "Công việc chính", "Đầu ra / gate"], roadmap,
              [1200, 1800, 4200, 2160], font_size=7.9)
    add_source_note(doc, "Dấu * là nhánh fine-tune, chỉ bắt đầu sau Gate 1. Nhân sự có thể chạy song song nếu đủ người.")

    add_heading(doc, "Nhân sự tối thiểu", 2)
    team = [
        ("Owner/PM", "0,2-0,4 FTE", "Scope, budget, gate, consent, stakeholder."),
        ("ML Engineer", "1,0 FTE trong POC/fine-tune", "Dataset, benchmark, training, error analysis, model card."),
        ("Backend/MLOps", "0,5-1,0 FTE", "ASR service, router, observability, deployment, on-call."),
        ("Ngôn ngữ/QA", "0,3-0,5 FTE", "Guideline transcript, adjudication, critical error review."),
        ("Annotator", "2-4 người part-time", "Nghe/sửa/QC 10.000 clip."),
        ("Privacy/Security", "0,1-0,2 FTE", "Consent, retention, access, incident, vendor review."),
    ]
    add_table(doc, ["Vai trò", "Mức tham gia", "Trách nhiệm"], team, [1900, 2100, 5360], font_size=8.5)


def add_quality_gates(doc: Document):
    add_heading(doc, "11. Tiêu chí nghiệm thu: thế nào là 'cao như Cloudflare'?", 1)
    add_body(doc, "Không thể cam kết tương đương trước khi có golden set. “Parity” phải được định nghĩa bằng ngưỡng định lượng, đo trên cùng audio, cùng normalization và theo từng nhóm vùng miền.")
    gates = [
        ("WER", "Candidate không tệ hơn Cloudflare quá 10% tương đối và không quá +1,0 điểm % tuyệt đối trên toàn bộ test."),
        ("CER", "Không tệ hơn +0,5 điểm % tuyệt đối; báo cáo riêng lỗi dấu/thanh."),
        ("Critical-token error", "Không cao hơn Cloudflare: phủ định, xưng hô, chủ thể, đồ vật, địa điểm, câu hỏi và từ khóa bài học."),
        ("Slice fairness", "Không slice vùng/tuổi/noise nào tệ hơn Cloudflare >2 điểm % WER nếu sample đủ; nếu thiếu sample phải ghi 'chưa kết luận'."),
        ("Latency", "P50 final ≤1,2 giây; P95 ≤2,5 giây sau khi người dùng dừng nói, phù hợp threshold hiện có của dự án."),
        ("Reliability", "Success ≥99,5%; fallback <10% ở canary và <5% sau tối ưu; không retry storm."),
        ("Cost", "TCO/1.000 phút thấp hơn Cloudflare ít nhất 30% ở lưu lượng dự báo sau khi tính ops, hoặc có lợi ích privacy/offline được phê duyệt."),
    ]
    add_table(doc, ["Metric", "Gate đề xuất"], gates, [2050, 7310], font_size=8.7)
    add_source_note(doc, "Các ngưỡng là đề xuất quản trị; cần điều chỉnh theo baseline thực tế và mức sai số thống kê của từng slice.")

    add_heading(doc, "Thiết kế benchmark", 2)
    for item in [
        "Golden test 1.500-2.000 clip, speaker-disjoint, khóa version; tối thiểu 200-300 clip cho mỗi nhóm vùng chính nếu có thể.",
        "Bootstrap confidence interval cho chênh lệch WER; không ra quyết định chỉ bằng một số trung bình.",
        "Báo cáo raw transcript và post-repair transcript riêng; manual review 100 lỗi nghiêm trọng nhất mỗi model.",
        "Load test ở concurrency 1, 5, 20 và peak dự báo; đo cold start, queue time, GPU utilization và OOM.",
        "Chạy tối thiểu 7 ngày shadow để bao phủ thiết bị/network/bối cảnh thực; không dùng output shadow làm nhãn tự động.",
    ]:
        add_list_item(doc, item)


def add_risks(doc: Document):
    add_heading(doc, "12. Rủi ro và biện pháp kiểm soát", 1)
    risks = [
        ("Dữ liệu trẻ em/giọng nói", "Cao", "Consent tách biệt; đại diện hợp pháp; data minimization; retention/xóa; mã hóa; hạn chế truy cập; tư vấn pháp lý."),
        ("Dataset lệch vùng", "Cao", "Quota và dashboard theo vùng; active collection; speaker-disjoint test; không suy diễn khi sample thấp."),
        ("Transcript giả làm ground truth", "Cao", "Human verification hai vòng; audit agreement; adjudication; không train trực tiếp từ output Cloudflare."),
        ("Hallucination/mất phủ định", "Cao", "Threshold, VAD, critical-token review, fallback, user retry; raw/postprocess metrics."),
        ("GPU nhàn rỗi", "Trung bình", "Serverless/scale-to-zero giai đoạn đầu; dedicated chỉ sau capacity test; batch clip ngắn."),
        ("Cold start / OOM", "Trung bình", "Min warm workers; VRAM headroom; request limits; queue; circuit breaker; Cloudflare fallback."),
        ("Model drift", "Trung bình", "Monthly golden regression; live review sample opt-in; alert theo slice; dataset/model versioning."),
        ("Lock-in tự xây", "Trung bình", "API adapter tương thích; OpenAI-style schema; container chuẩn; model registry; tài liệu runbook."),
    ]
    add_table(doc, ["Rủi ro", "Mức", "Kiểm soát"], risks, [2000, 900, 6460], font_size=8.3,
              alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])

    add_heading(doc, "Bảo vệ dữ liệu", 2)
    add_body(
        doc,
        "Luật Bảo vệ dữ liệu cá nhân số 91/2025/QH15 có hiệu lực 01/01/2026; pháp luật có quy định riêng đối với dữ liệu cá nhân của trẻ em và quyền của người đại diện. Nghị định 356/2025/NĐ-CP hướng dẫn thi hành cũng có hiệu lực từ 01/01/2026. Báo cáo này không phải ý kiến pháp lý; dự án cần privacy/legal review trước khi lưu audio để train. [14][15]",
    )
    add_body(
        doc,
        "Cloudflare cho biết Customer Content thuộc trách nhiệm của khách hàng, không được chia sẻ cho khách hàng khác và không được dùng để train/cải thiện model nếu không có đồng ý rõ; nội dung có thể được lưu nếu người dùng chủ động kết hợp các dịch vụ storage. Điều này không thay thế nghĩa vụ consent/retention của ứng dụng. [4]",
    )


def add_action_plan(doc: Document):
    add_heading(doc, "13. Checklist hành động 30 ngày", 1)
    actions = [
        ("Ngày 1-3", "Xuất dashboard Cloudflare theo model; đối soát invoice; đo phút audio; xác định ASR/LLM/TTS share.", "Backend + Owner"),
        ("Ngày 1-5", "Chốt mục tiêu user, tuổi, vùng, clip duration, traffic hiện tại/12 tháng; quyết định có trẻ em hay không.", "Owner + Product"),
        ("Ngày 4-8", "Thiết kế consent/retention/delete và data schema; legal/privacy review.", "Privacy + Backend"),
        ("Ngày 6-12", "Viết transcript guideline; tuyển annotator; chọn 1.500 clip stratified; loại PII không cần thiết.", "ML + Language QA"),
        ("Ngày 10-18", "Gán nhãn 2 vòng; tính agreement; khóa golden v1.", "Annotator + QA"),
        ("Ngày 15-22", "Chạy Cloudflare baseline; log raw/post-repair; report WER/CER/slice/latency.", "ML + Backend"),
        ("Ngày 20-27", "POC Whisper Turbo và PhoWhisper-medium trên GPU thuê theo giờ; benchmark single/concurrency.", "ML + MLOps"),
        ("Ngày 28-30", "Họp Gate 1: dừng, hybrid hay fine-tune; phê duyệt ngân sách giai đoạn sau.", "Owner + Stakeholders"),
    ]
    add_table(doc, ["Thời gian", "Việc cần làm", "Owner"], actions, [1300, 6100, 1960], font_size=8.3)
    add_callout(
        doc,
        "Gate 1 - điều kiện tiếp tục self-host",
        "Chỉ tiếp tục nếu (a) model local đạt quality gate hoặc có đường cải thiện rõ, và (b) forecast/TCO hoặc yêu cầu privacy/offline biện minh cho chi phí dự án. Nếu không, dừng ở Cloudflare + tối ưu managed là quyết định đúng.",
        kind="positive",
    )


def add_sources(doc: Document):
    add_heading(doc, "14. Nguồn tham khảo và giả định", 1)
    add_body(doc, "Nguồn trực tuyến được kiểm tra ngày 27/07/2026. Giá, model và giới hạn có thể thay đổi; cần kiểm tra lại trước quyết định mua/triển khai.")
    sources_num_id = add_numbering_definition(doc, bullet=False)
    sources = [
        ("Cloudflare - whisper-large-v3-turbo model page (0,00051 USD/phút; tham số ASR)", "https://developers.cloudflare.com/workers-ai/models/whisper-large-v3-turbo/"),
        ("Cloudflare Workers AI Pricing (Neurons, free allocation, giá audio/text/TTS)", "https://developers.cloudflare.com/workers-ai/platform/pricing/"),
        ("Cloudflare Workers AI Limits (ASR 720 requests/phút)", "https://developers.cloudflare.com/workers-ai/platform/limits/"),
        ("Cloudflare - Your Data and Workers AI", "https://developers.cloudflare.com/workers-ai/platform/data-usage/"),
        ("Cloudflare Workers Pricing (Workers Paid tối thiểu 5 USD/tháng)", "https://developers.cloudflare.com/workers/platform/pricing/"),
        ("PhoWhisper paper / model: 843,79 giờ, đa giọng Việt, WER benchmark", "https://arxiv.org/abs/2406.02555"),
        ("Hugging Face - Whisper docs và fine-tuning multilingual ASR", "https://huggingface.co/docs/transformers/model_doc/whisper"),
        ("SYSTRAN faster-whisper - CTranslate2, INT8 và benchmark", "https://github.com/SYSTRAN/faster-whisper"),
        ("RunPod GPU pricing (cập nhật 17/07/2026)", "https://www.runpod.io/pricing"),
        ("OpenAI Whisper large-v3-turbo model card và MIT license", "https://huggingface.co/openai/whisper-large-v3-turbo"),
        ("Mozilla Common Voice Vietnamese v23 (22 giờ, 361 speakers, CC0)", "https://dev.datacollective.mozillafoundation.org/datasets/cmflnn484oaiy1nwkq2dp76ig"),
        ("VIVOS Vietnamese Speech Corpus - license CC BY-NC-SA 4.0", "https://live.european-language-grid.eu/catalogue/corpus/22131"),
        ("whisper.cpp - Android/iOS/WebAssembly, quantization, CPU/GPU", "https://github.com/ggml-org/whisper.cpp"),
        ("Luật số 91/2025/QH15 - Bảo vệ dữ liệu cá nhân", "https://congbao.chinhphu.vn/van-ban/luat-so-91-2025-qh15-45578/57730.htm"),
        ("Nghị định 356/2025/NĐ-CP hướng dẫn Luật Bảo vệ dữ liệu cá nhân", "https://vbpl.vn/TW/Pages/vbpq-toanvan.aspx?ItemID=187276"),
        ("VinAI PhoWhisper-small model card, file size và BSD-3-Clause license", "https://huggingface.co/vinai/PhoWhisper-small"),
        ("VinAI PhoWhisper-medium model card, file size và BSD-3-Clause license", "https://huggingface.co/vinai/PhoWhisper-medium"),
        ("Hugging Face PEFT - LoRA và parameter-efficient fine-tuning", "https://huggingface.co/docs/peft/en/index"),
        ("RunPod Serverless pricing - GPU 16 GB và 24 GB tính theo giây", "https://docs.runpod.io/serverless/pricing"),
    ]
    for idx, (title, url) in enumerate(sources, start=1):
        p = doc.add_paragraph(style="Normal")
        apply_num(p, sources_num_id)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(title + ". ")
        set_run_font(r, size=9.5)
        add_hyperlink(p, url, url)

    add_heading(doc, "Giả định tính toán", 2)
    assumptions = [
        "Giá ASR dùng 0,00051 USD/phút từ model page (thận trọng hơn số 0,0005 USD/phút ở bảng platform).",
        "Một tháng = 720 giờ; Cloudflare free allocation tính 30 ngày khi minh họa tháng.",
        "Tỷ giá kế hoạch 1 USD = 26.000 VND, không phải báo giá ngoại hối.",
        "Điểm hòa vốn GPU compute-only không gồm CPU/RAM/storage/egress/ops/cold start; vùng TCO 250-350 USD/tháng là planning estimate.",
        "Chi phí nhân sự và gán nhãn là khoảng lập ngân sách, không phải báo giá nhà cung cấp.",
        "Không có 10.000 audio trong workspace được audit; các kịch bản tổng giờ dựa trên độ dài trung bình giả định.",
    ]
    for item in assumptions:
        add_list_item(doc, item)


def audit_document(doc: Document):
    # Basic structural QA before save.
    assert len(doc.paragraphs) > 100
    forbidden = ("turn0search", "turn1search", "cite", "TODO", "PLACEHOLDER")
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    for token in forbidden:
        assert token not in full_text, f"Forbidden token: {token}"
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        assert tbl_pr.find(qn("w:tblW")) is not None
        assert tbl_pr.find(qn("w:tblInd")) is not None
        assert table._tbl.tblGrid is not None


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    make_architecture_diagram(ARCH_PATH)
    make_cost_chart(COST_PATH)

    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    doc._bullet_num_id = add_numbering_definition(doc, bullet=True)
    doc._decimal_num_id = add_numbering_definition(doc, bullet=False)
    props = doc.core_properties
    props.title = "Báo cáo phương án AI nhận dạng giọng nói tiếng Việt"
    props.subject = "Tối ưu chi phí Cloudflare ASR và lộ trình self-host/fine-tune"
    props.author = "Nhóm dự án"
    props.keywords = "ASR, tiếng Việt, Cloudflare, Whisper, PhoWhisper, fine-tune, chi phí"
    props.comments = "Bản nghiên cứu kỹ thuật ngày 27/07/2026"

    add_cover(doc)
    add_executive_summary(doc)
    add_current_state(doc)
    add_cost_analysis(doc)
    add_data_analysis(doc)
    add_options(doc)
    add_architecture(doc)
    add_training_guide(doc)
    add_model_playbooks(doc)
    add_budget(doc)
    add_timeline(doc)
    add_quality_gates(doc)
    add_risks(doc)
    add_action_plan(doc)
    add_sources(doc)

    configure_sections(doc)
    audit_document(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
