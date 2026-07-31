from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from lxml import etree
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import build_asr_report as base


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "reports" / "Build AI - ban hoan thien.docx"
OUTPUT = ROOT / "reports" / "Build AI - chinh sua trinh bay.docx"

CONTENT_PART_PREFIXES = (
    "word/document.xml",
    "word/header",
    "word/footer",
    "word/footnotes.xml",
    "word/endnotes.xml",
    "word/comments.xml",
)


def content_signature(path: Path):
    values = []
    hyperlinks = []
    with ZipFile(path) as package:
        for name in sorted(package.namelist()):
            if not name.endswith(".xml") or not name.startswith(CONTENT_PART_PREFIXES):
                continue
            root = etree.fromstring(package.read(name))
            for node in root.iter():
                local = etree.QName(node).localname
                if local in {"t", "instrText", "delText", "tab", "br"}:
                    values.append((name, local, node.text or ""))

        rel_name = "word/_rels/document.xml.rels"
        if rel_name in package.namelist():
            rel_root = etree.fromstring(package.read(rel_name))
            for rel in rel_root:
                rel_type = rel.get("Type", "")
                if rel_type.endswith("/hyperlink"):
                    hyperlinks.append((rel.get("Id"), rel.get("Target"), rel.get("TargetMode")))
    return values, sorted(hyperlinks)


def set_style_font(style, name: str, size: float, color: str, *, bold=None, italic=None):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = base.rgb(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def format_styles(doc: Document):
    normal = doc.styles["Normal"]
    set_style_font(normal, "Calibri", 11, base.BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_specs = {
        "Heading 1": (16, base.BLUE, 18, 8),
        "Heading 2": (13, base.BLUE, 12, 6),
        "Heading 3": (12, base.DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        set_style_font(style, "Calibri", size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Kicker" in doc.styles:
        style = doc.styles["Kicker"]
        set_style_font(style, "Calibri", 10, base.GOLD, bold=True)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(7)

    if "Report Title" in doc.styles:
        style = doc.styles["Report Title"]
        set_style_font(style, "Calibri", 28, base.NAVY, bold=True)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    if "Report Subtitle" in doc.styles:
        style = doc.styles["Report Subtitle"]
        set_style_font(style, "Calibri", 14, base.DARK_BLUE)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(16)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.keep_with_next = True

    if "Table Citation" in doc.styles:
        style = doc.styles["Table Citation"]
        set_style_font(style, "Calibri", 8.5, base.MUTED, italic=True)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.keep_with_next = False


def has_numbering(paragraph) -> bool:
    ppr = paragraph._p.pPr
    return ppr is not None and ppr.numPr is not None


def format_paragraphs(doc: Document):
    title_seen = False
    first_heading_seen = False
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        fmt = paragraph.paragraph_format

        if style_name == "Report Title":
            title_seen = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif style_name == "Report Subtitle":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif style_name == "Kicker":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif style_name.startswith("Heading "):
            first_heading_seen = True
            fmt.keep_with_next = True
            fmt.widow_control = True
        elif style_name == "Table Citation":
            fmt.keep_with_next = False
        else:
            fmt.widow_control = True
            if has_numbering(paragraph):
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(5)
                fmt.line_spacing = 1.12
            elif paragraph.text.strip():
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(6)
                fmt.line_spacing = 1.10

        # The short metadata block below the title receives tighter rhythm only.
        if title_seen and not first_heading_seen and style_name == "Normal" and paragraph.text.strip():
            fmt.space_after = Pt(3)
            fmt.line_spacing = 1.05

        for run in paragraph.runs:
            base.set_run_font(run, name="Calibri")


def set_row_no_split(row):
    trpr = row._tr.get_or_add_trPr()
    if trpr.find(qn("w:cantSplit")) is None:
        trpr.append(OxmlElement("w:cantSplit"))


def get_grid_widths(table):
    grid = table._tbl.tblGrid
    widths = []
    if grid is not None:
        for col in grid.findall(qn("w:gridCol")):
            value = col.get(qn("w:w"))
            if value:
                widths.append(int(value))
    if len(widths) != len(table.columns):
        widths = [9360 // len(table.columns)] * len(table.columns)
        widths[-1] += 9360 - sum(widths)
    return widths


def format_table_text(cell, *, header=False, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, color=base.BLACK, size=9.0):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        paragraph.paragraph_format.widow_control = True
        for run in paragraph.runs:
            base.set_run_font(
                run,
                name="Calibri",
                size=size,
                color=color,
                bold=True if header or bold else run.bold,
            )


def format_tables(doc: Document):
    for table in doc.tables:
        widths = get_grid_widths(table)
        base.set_table_geometry(table, widths, indent_dxa=120)
        base.set_table_borders(table, color="C9D2DC", size="5")
        table.alignment = 0
        table.autofit = False

        is_callout = len(table.columns) == 1 and len(table.rows) == 1
        if is_callout:
            base.set_repeat_table_header(table.rows[0])
            cell = table.cell(0, 0)
            base.set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
            format_table_text(cell, size=10.2, color=base.INK)
            set_row_no_split(table.rows[0])
            continue

        base.set_repeat_table_header(table.rows[0])
        for col_index, cell in enumerate(table.rows[0].cells):
            base.set_cell_shading(cell, "E8EEF5")
            header_align = WD_ALIGN_PARAGRAPH.LEFT if col_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            format_table_text(cell, header=True, align=header_align, color=base.NAVY, size=9.0)
        set_row_no_split(table.rows[0])

        for row_index, row in enumerate(table.rows[1:], start=1):
            set_row_no_split(row)
            row_text = " ".join(cell.text.strip() for cell in row.cells).lower()
            is_total = "tổng" in row_text or "quy đổi" in row_text
            for col_index, cell in enumerate(row.cells):
                if is_total:
                    base.set_cell_shading(cell, "EAF2F8")
                elif row_index % 2 == 0:
                    base.set_cell_shading(cell, "FAFBFC")

                if len(row.cells) == 2:
                    align = WD_ALIGN_PARAGRAPH.LEFT if col_index == 0 else WD_ALIGN_PARAGRAPH.RIGHT
                elif col_index == 0:
                    align = WD_ALIGN_PARAGRAPH.LEFT
                elif any(token in table.rows[0].cells[col_index].text.lower() for token in ("chi phí", "ngân sách", "usd", "vnđ")):
                    align = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    align = WD_ALIGN_PARAGRAPH.LEFT
                format_table_text(cell, align=align, bold=is_total, color=base.INK if is_total else base.BLACK, size=8.8)

        base.set_table_geometry(table, widths, indent_dxa=120)


def format_sections(doc: Document):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.header_distance = Inches(0.40)
        section.footer_distance = Inches(0.40)

        for paragraph in section.header.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                base.set_run_font(run, name="Calibri", size=8, color=base.MUTED, bold=True)

        for paragraph in section.first_page_header.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                base.set_run_font(run, name="Calibri", size=8, color=base.MUTED, bold=True)

        for footer in (section.footer, section.first_page_footer):
            for paragraph in footer.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    base.set_run_font(run, name="Calibri", size=8.5, color=base.MUTED)


def structural_counts(doc: Document):
    return (
        len(doc.paragraphs),
        len(doc.tables),
        tuple(len(table.rows) for table in doc.tables),
        tuple(len(table.columns) for table in doc.tables),
    )


def main():
    before_signature = content_signature(SOURCE)
    source_doc = Document(SOURCE)
    before_counts = structural_counts(source_doc)

    format_styles(source_doc)
    format_paragraphs(source_doc)
    format_tables(source_doc)
    format_sections(source_doc)
    source_doc.save(OUTPUT)

    after_doc = Document(OUTPUT)
    after_counts = structural_counts(after_doc)
    after_signature = content_signature(OUTPUT)

    assert before_counts == after_counts, (before_counts, after_counts)
    assert before_signature == after_signature, "Text, field, break, tab, or hyperlink content changed"
    print(f"CONTENT_UNCHANGED=YES counts={before_counts}")
    print(OUTPUT)


if __name__ == "__main__":
    main()
