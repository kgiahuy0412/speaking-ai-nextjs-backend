from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import build_asr_report as base


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
OUT_PATH = OUT_DIR / "Build AI - ban hoan thien.docx"

USD_TO_VND = 26_000


def usd(value: float) -> str:
    return f"{value:,.2f} USD".replace(",", " ")


def vnd(value: float) -> str:
    return f"{value * USD_TO_VND:,.0f} VNĐ".replace(",", ".")


def set_section_chrome(doc: Document):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        hp = section.header.paragraphs[0]
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        base.set_run_font(
            hp.add_run("BÁO CÁO PHƯƠNG ÁN NHẬN DẠNG GIỌNG NÓI TIẾNG VIỆT"),
            size=8.5,
            color=base.MUTED,
            bold=True,
        )

        fp = section.first_page_header.paragraphs[0]
        fp.clear()
        base.set_run_font(
            fp.add_run("TÀI LIỆU PHÂN TÍCH KỸ THUẬT VÀ CHI PHÍ"),
            size=8.5,
            color=base.MUTED,
            bold=True,
        )

        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        base.set_run_font(footer.add_run("Trang "), size=9, color=base.MUTED)
        base.add_page_field(footer, "PAGE")
        base.set_run_font(footer.add_run(" / "), size=9, color=base.MUTED)
        base.add_page_field(footer, "NUMPAGES")

        first_footer = section.first_page_footer.paragraphs[0]
        first_footer.clear()
        first_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        base.set_run_font(
            first_footer.add_run("Bản hoàn thiện từ tài liệu Build AI - tỷ giá kế hoạch 1 USD = 26.000 VNĐ"),
            size=8.5,
            color=base.MUTED,
            italic=True,
        )


def add_body(doc: Document, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        base.set_run_font(p.add_run(bold_lead), bold=True, color=base.INK)
        base.set_run_font(p.add_run(text[len(bold_lead):]), color=base.BLACK)
    else:
        base.set_run_font(p.add_run(text), color=base.BLACK)
    return p


def add_callout(doc: Document, label: str, text: str, *, kind: str):
    table = base.add_callout(doc, label, text, kind=kind)
    base.set_table_geometry(table, [9360], indent_dxa=120)
    base.set_repeat_table_header(table.rows[0])
    return table


def add_source(doc: Document, title: str, url: str):
    p = doc.add_paragraph(style="Normal")
    base.apply_num(p, doc._decimal_num_id)
    p.paragraph_format.space_after = Pt(6)
    base.set_run_font(p.add_run(f"{title}: "), size=9.5)
    base.add_hyperlink(p, "Mở tài liệu chính thức", url)


def add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph(style="Kicker")
    base.set_run_font(p.add_run("BÁO CÁO ĐỀ XUẤT"), size=10, color=base.GOLD, bold=True)

    p = doc.add_paragraph(style="Report Title")
    p.paragraph_format.space_after = Pt(6)
    base.set_run_font(
        p.add_run("Phương án xây dựng AI nhận dạng giọng nói tiếng Việt"),
        size=27,
        color=base.NAVY,
        bold=True,
    )

    p = doc.add_paragraph(style="Report Subtitle")
    base.set_run_font(
        p.add_run("So sánh Cloudflare Workers AI và fine-tune PhoWhisper-small trên GPU thuê"),
        size=14,
        color=base.DARK_BLUE,
    )

    metadata = [
        ("Quy mô", "10.000 người dùng"),
        ("Kịch bản chính", "Mỗi người dùng tổng cộng 10 phút trong một tháng"),
        ("Kịch bản tải cao", "Mỗi người dùng 10 phút mỗi ngày trong 30 ngày"),
        ("Tỷ giá kế hoạch", "1 USD = 26.000 VNĐ"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        base.set_run_font(p.add_run(f"{label}: "), size=10.5, color=base.INK, bold=True)
        base.set_run_font(p.add_run(value), size=10.5, color=base.INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_callout(
        doc,
        "Kết luận điều hành",
        "Ở mức 100.000 phút/tháng, Cloudflare có chi phí trực tiếp thấp hơn hạ tầng GPU chạy 24/7. Fine-tune PhoWhisper-small vẫn đáng thực hiện như một POC chất lượng, nhưng chỉ nên chuyển production sau khi benchmark độ chính xác, tốc độ và khả năng chịu tải.",
        kind="positive",
    )


def add_overview(doc: Document):
    base.add_heading(doc, "1. Phạm vi và giả định", 1)
    for item in [
        "Phương án 1 sử dụng Cloudflare Workers Paid và model Whisper Large V3 Turbo.",
        "Phương án 2 fine-tune vinai/PhoWhisper-small trên GPU RTX A5000 24 GB, sau đó tự vận hành model.",
        "Giá Cloudflare được tính theo phút audio; giá GPU được tính theo số giờ máy thực chạy.",
        "Chi phí nhân sự phát triển, gán nhãn transcript, thuế và phí thanh toán quốc tế chưa được đưa vào tổng ngân sách.",
        "Các mức giá là dữ liệu tham khảo tại thời điểm lập báo cáo và cần được kiểm tra lại trước khi mua dịch vụ.",
    ]:
        base.add_list_item(doc, item)

    base.add_heading(doc, "2. Tóm tắt chi phí", 1)
    rows = [
        ("Cloudflare - 100.000 phút/tháng", "Hằng tháng", "1.370.200 VNĐ", "Bao gồm Workers Paid và quota miễn phí 30 ngày"),
        ("Cloudflare - 3.000.000 phút/tháng", "Hằng tháng", "39.824.200 VNĐ", "Kịch bản 10 phút/người/ngày"),
        ("Fine-tune PhoWhisper-small", "Một lần/train lại", "390.000-910.000 VNĐ", "GPU train; chưa gồm transcript và nhân sự"),
        ("RTX A5000 chạy 24/7", "Hằng tháng", "5.054.400 VNĐ", "Chỉ riêng GPU"),
        ("Toàn bộ hạ tầng self-host", "Hằng tháng", "5,3-6,1 triệu VNĐ", "GPU, storage, API, queue, log và backup"),
    ]
    base.add_table(
        doc,
        ["Hạng mục", "Loại chi phí", "Ngân sách", "Ghi chú"],
        rows,
        [2700, 1700, 2100, 2860],
        font_size=8.5,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    base.add_source_note(doc, "Bảng 1. Tổng hợp nhanh; số liệu chi tiết được trình bày tại các phần tiếp theo.")
    doc.add_page_break()


def add_cloudflare(doc: Document):
    base.add_heading(doc, "3. Phương án 1 - Cloudflare Workers Paid + Whisper", 1)
    add_body(doc, "Cloudflare vận hành hạ tầng GPU và tính phí theo số phút audio. Dự án không cần tự quản lý model, CUDA, hàng đợi GPU hoặc khả năng mở rộng inference.")

    base.add_heading(doc, "3.1. Kịch bản 10 phút/người/tháng", 2)
    add_body(doc, "Tổng khối lượng: 10.000 người × 10 phút = 100.000 phút/tháng.", bold_lead="Tổng khối lượng:")
    rows = [
        ("Whisper trước quota miễn phí", "51,00 USD"),
        ("Trừ quota miễn phí 30 ngày", "-3,30 USD"),
        ("Workers Paid", "+5,00 USD"),
        ("Tổng", "52,70 USD/tháng"),
        ("Quy đổi 26.000 VNĐ/USD", "1.370.200 VNĐ/tháng"),
    ]
    base.add_table(doc, ["Khoản phí", "Chi phí"], rows, [6200, 3160], font_size=9.5, alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT])
    base.add_source_note(doc, "Bảng 2. Chi phí Cloudflare cho 100.000 phút audio phân bổ trong 30 ngày.")

    base.add_heading(doc, "3.2. Kịch bản 10 phút/người/ngày", 2)
    add_body(doc, "Tổng khối lượng: 10.000 người × 10 phút × 30 ngày = 3.000.000 phút/tháng.", bold_lead="Tổng khối lượng:")
    rows = [
        ("Whisper trước quota miễn phí", "1.530,00 USD"),
        ("Trừ quota miễn phí 30 ngày", "-3,30 USD"),
        ("Workers Paid", "+5,00 USD"),
        ("Tổng", "1.531,70 USD/tháng"),
        ("Quy đổi 26.000 VNĐ/USD", "39.824.200 VNĐ/tháng"),
    ]
    base.add_table(doc, ["Khoản phí", "Chi phí"], rows, [6200, 3160], font_size=9.5, alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT])
    base.add_source_note(doc, "Bảng 3. Chi phí Cloudflare khi mỗi người dùng phát sinh 10 phút mỗi ngày.")

    base.add_heading(doc, "3.3. Ưu điểm và giới hạn", 2)
    for item in [
        "Triển khai nhanh; dự án hiện đã có lớp tích hợp Cloudflare Workers AI.",
        "Không cần thuê và vận hành GPU riêng.",
        "Chi phí tăng tuyến tính theo tổng số phút audio.",
        "Không thể fine-tune model Cloudflare bằng bộ 10.000 audio của dự án.",
        "Cần kiểm soát fallback để tránh phát sinh phí ngoài dự kiến khi Cloudflare lỗi hoặc hết quota.",
    ]:
        base.add_list_item(doc, item)
    doc.add_page_break()


def add_self_host(doc: Document):
    base.add_heading(doc, "4. Phương án 2 - Fine-tune PhoWhisper-small", 1)
    add_body(doc, "PhoWhisper-small là model nhận dạng giọng nói đã được tối ưu cho tiếng Việt. Phương án đề xuất là fine-tune bằng LoRA trên dữ liệu của dự án, không huấn luyện model từ đầu.")

    base.add_heading(doc, "4.1. Cấu hình GPU và chi phí huấn luyện", 2)
    rows = [
        ("Model", "vinai/PhoWhisper-small"),
        ("GPU", "RTX A5000"),
        ("VRAM", "24 GB"),
        ("Giá RunPod Pods", "0,27 USD/giờ"),
        ("Thời gian một lần train", "15-30 giờ"),
        ("Ba lần thử nghiệm", "45-90 giờ"),
    ]
    base.add_table(doc, ["Thông số", "Giá trị"], rows, [3600, 5760], font_size=9.5)
    base.add_source_note(doc, "Bảng 4. Cấu hình lập ngân sách cho fine-tune PhoWhisper-small.")

    add_body(doc, "Chi phí GPU thuần cho ba lần thử nghiệm:", bold_lead="Chi phí GPU thuần")
    for item in [
        "Tối thiểu: 45 giờ × 0,27 USD = 12,15 USD, tương đương 315.900 VNĐ.",
        "Tối đa: 90 giờ × 0,27 USD = 24,30 USD, tương đương 631.800 VNĐ.",
        "Ngân sách nên dự phòng: 15-35 USD, tương đương khoảng 390.000-910.000 VNĐ.",
    ]:
        base.add_list_item(doc, item)

    add_callout(
        doc,
        "Không bao gồm",
        "Ngân sách train chưa gồm công việc nghe và sửa transcript, làm sạch dữ liệu, phát triển pipeline, đánh giá WER/CER và thời gian của kỹ sư.",
        kind="caution",
    )

    base.add_heading(doc, "4.2. Chi phí triển khai model sau huấn luyện", 2)
    add_body(doc, "Nếu giữ RTX A5000 hoạt động liên tục 24 giờ/ngày trong 30 ngày:")
    add_body(doc, "0,27 USD × 24 giờ × 30 ngày = 194,40 USD/tháng = 5.054.400 VNĐ/tháng.")
    rows = [
        ("GPU RTX A5000 chạy 24/7", "194,40 USD", "5.054.400 VNĐ"),
        ("Storage/checkpoint", "3-10 USD", "78.000-260.000 VNĐ"),
        ("API server và queue", "5-20 USD", "130.000-520.000 VNĐ"),
        ("Log, backup, monitoring", "2-10 USD", "52.000-260.000 VNĐ"),
        ("Tổng hạ tầng dự kiến", "204-234 USD", "Khoảng 5,3-6,1 triệu VNĐ"),
    ]
    base.add_table(doc, ["Khoản vận hành", "USD/tháng", "VNĐ/tháng"], rows, [4300, 2200, 2860], font_size=8.8, alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT])
    base.add_source_note(doc, "Bảng 5. Chi phí self-host ước tính khi duy trì một GPU liên tục; chưa gồm phương án dự phòng nhiều GPU.")

    base.add_heading(doc, "4.3. Các điều kiện bắt buộc trước production", 2)
    for item in [
        "Transcript phải được con người kiểm tra; dữ liệu sai sẽ làm model học sai.",
        "Chia train/validation/test theo người nói, không để cùng một người xuất hiện ở cả train và test.",
        "Đánh giá riêng theo giọng Bắc, Trung, Nam, độ tuổi, thiết bị và mức nhiễu.",
        "Load test throughput và concurrency; một RTX A5000 không mặc nhiên đủ cho mọi kịch bản 10.000 người dùng.",
        "Chuẩn bị queue, retry, timeout, autoscaling, monitoring và phương án fallback.",
    ]:
        base.add_list_item(doc, item)
    doc.add_page_break()


def add_comparison(doc: Document):
    base.add_heading(doc, "5. So sánh và khuyến nghị", 1)
    rows = [
        ("Chi phí khởi tạo", "Thấp", "390.000-910.000 VNĐ tiền GPU train; chưa gồm dữ liệu và nhân sự"),
        ("Chi phí ở 100.000 phút/tháng", "1.370.200 VNĐ/tháng", "5,3-6,1 triệu VNĐ/tháng nếu GPU chạy 24/7"),
        ("Khả năng fine-tune", "Không", "Có"),
        ("Thời gian triển khai", "Nhanh", "Nhiều tuần, phụ thuộc dữ liệu"),
        ("Vận hành GPU", "Cloudflare chịu trách nhiệm", "Dự án tự chịu trách nhiệm"),
        ("Rủi ro mở rộng", "Thấp hơn", "Cần benchmark, queue và autoscaling"),
    ]
    base.add_table(doc, ["Tiêu chí", "Cloudflare", "PhoWhisper-small self-host"], rows, [2350, 2700, 4310], font_size=8.5)
    base.add_source_note(doc, "Bảng 6. So sánh theo cấu hình và giả định trong báo cáo.")

    base.add_heading(doc, "5.1. Khuyến nghị theo giai đoạn", 2)
    for item in [
        "Giai đoạn hiện tại: giữ Cloudflare cho production để đảm bảo độ ổn định và chi phí thấp ở mức 100.000 phút/tháng.",
        "Song song: thực hiện POC PhoWhisper-small trên 1.000 audio trước, sau đó mới mở rộng lên 10.000 audio.",
        "Chỉ chuyển sang self-host khi model đạt yêu cầu chất lượng và load test chứng minh chi phí trên mỗi phút thấp hơn phương án API.",
        "Ở kịch bản 3.000.000 phút/tháng, cần benchmark số GPU thực tế; không được so trực tiếp 39,82 triệu VNĐ với chi phí của một GPU nếu chưa chứng minh đủ công suất.",
    ]:
        base.add_list_item(doc, item)

    add_callout(
        doc,
        "Quyết định đề xuất",
        "Cloudflare là phương án production ngắn hạn. PhoWhisper-small là hướng R&D trung hạn để sở hữu model và tối ưu cho giọng nói thực tế của người dùng Việt Nam.",
        kind="positive",
    )

    base.add_heading(doc, "6. Lộ trình POC PhoWhisper-small", 1)
    rows = [
        ("1", "Audit dữ liệu và quyền sử dụng", "2-3 ngày"),
        ("2", "Chuẩn hóa audio và transcript", "1-4 tuần"),
        ("3", "Chạy baseline trên tập test", "2-3 ngày"),
        ("4", "Fine-tune LoRA và thử nghiệm", "3-7 ngày"),
        ("5", "Đánh giá vùng miền và lỗi từ quan trọng", "3-5 ngày"),
        ("6", "Xây API, queue và load test", "1-2 tuần"),
    ]
    base.add_table(doc, ["Bước", "Công việc", "Thời gian"], rows, [900, 6500, 1960], font_size=9.0, alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])
    base.add_source_note(doc, "Bảng 7. Lộ trình tham khảo; phần transcript thường chiếm nhiều thời gian nhất.")


def add_sources(doc: Document):
    base.add_heading(doc, "7. Nguồn tham khảo", 1)
    add_source(doc, "Cloudflare - Whisper Large V3 Turbo", "https://developers.cloudflare.com/workers-ai/models/whisper-large-v3-turbo/")
    add_source(doc, "Cloudflare - Workers AI pricing", "https://developers.cloudflare.com/workers-ai/platform/pricing/")
    add_source(doc, "Cloudflare - Workers Paid pricing", "https://developers.cloudflare.com/workers/platform/pricing/")
    add_source(doc, "RunPod - GPU và storage pricing", "https://www.runpod.io/pricing")
    add_source(doc, "VinAI - PhoWhisper", "https://github.com/VinAIResearch/PhoWhisper")

    base.add_heading(doc, "Ghi chú", 2)
    add_body(doc, "Tài liệu này phục vụ so sánh kỹ thuật và lập ngân sách sơ bộ. Giá dịch vụ, tỷ giá, thời gian train và hiệu năng GPU có thể thay đổi; cần chạy POC trước khi đưa ra cam kết production.")


def audit(doc: Document):
    assert len(doc.paragraphs) >= 55
    assert len(doc.tables) >= 10
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    for token in ("TODO", "PLACEHOLDER", "turn27", "turn28"):
        assert token not in full_text


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    base.configure_styles(doc)
    doc._bullet_num_id = base.add_numbering_definition(doc, bullet=True)
    doc._decimal_num_id = base.add_numbering_definition(doc, bullet=False)
    set_section_chrome(doc)

    props = doc.core_properties
    props.title = "Phương án xây dựng AI nhận dạng giọng nói tiếng Việt"
    props.subject = "So sánh Cloudflare Workers AI và fine-tune PhoWhisper-small"
    props.author = "Nhóm dự án"
    props.keywords = "Cloudflare, Whisper, PhoWhisper, GPU, RunPod, ASR, tiếng Việt"
    props.comments = "Bản hoàn thiện từ tài liệu Build AI"

    add_title_block(doc)
    add_overview(doc)
    add_cloudflare(doc)
    add_self_host(doc)
    add_comparison(doc)
    add_sources(doc)

    audit(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
